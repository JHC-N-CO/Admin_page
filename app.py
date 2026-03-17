from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file, send_from_directory, after_this_request, make_response
import os
import pandas as pd
from werkzeug.utils import secure_filename
from flask_mail import Mail, Message
import unicodedata
import logging
from datetime import datetime, timedelta
from io import BytesIO
import qrcode
from urllib.parse import urlencode
from urllib.parse import quote, unquote
import uuid
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
import re
from dateutil import parser
from werkzeug.security import generate_password_hash, check_password_hash
from zipfile import ZipFile
import tempfile
import shutil
import zipfile

# SQLAlchemy 및 모델 import
from flask_sqlalchemy import SQLAlchemy
from models import db, Event, Participant, ParticipantFile, User, UserSession, Member, MemberDisplaySettings
from config import config

app = Flask(__name__, static_folder='static')

# 환경 설정
config_name = os.environ.get('FLASK_ENV', 'development')
app.config.from_object(config[config_name])

# 이메일 설정 디버깅 (실제 설정값 확인)
logging.info(f"Email Config - MAIL_SERVER: {app.config.get('MAIL_SERVER')}")
logging.info(f"Email Config - MAIL_USERNAME: {app.config.get('MAIL_USERNAME')}")
logging.info(f"Email Config - MAIL_DEFAULT_SENDER: {app.config.get('MAIL_DEFAULT_SENDER')}")

# SQLAlchemy 초기화
db.init_app(app)

# 로깅 설정
logging.basicConfig(level=logging.DEBUG)

# QR 코드 저장 폴더 설정 (다운로드 폴더로 변경됨)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def generate_qr(number):
    """QR 코드 생성 및 사용자 PC 다운로드 폴더에 저장 (행사명/이벤트ID 폴더 없이)"""
    try:
        qr = qrcode.make(str(number))
        qr_filename = f"{number}.png"
        # 환경에 따른 다운로드 폴더 설정
        if os.path.exists("/app/downloads"):  # Docker 환경 (서버)
            download_base = "/app/downloads"
        else:  # 로컬 환경
            home_dir = os.path.expanduser("~")
            download_base = os.path.join(home_dir, "Downloads")
        download_folder = os.path.join(download_base, "QR_Codes")
        os.makedirs(download_folder, exist_ok=True)
        qr_path = os.path.join(download_folder, qr_filename)
        qr.save(qr_path)
        # 전체 경로 반환
        return qr_path
    except Exception as e:
        logging.error(f"QR generation failed for {number}: {str(e)}")
        raise

# 업로드 폴더 설정
app.config['UPLOAD_FOLDER'] = 'uploads'
SUBFOLDERS = {
    'cv': 'CV',
    'photo': 'Photo',
    'ppt': 'PPT',
    'script': 'Script'
}
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
for subfolder in SUBFOLDERS.values():
    os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], subfolder), exist_ok=True)
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'pdf', 'png', 'jpg', 'jpeg', 'docx', 'ppt', 'pptx'}
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg'}

# 이미지 업로드를 위한 설정 추가
app.config['IMAGE_UPLOAD_FOLDER'] = 'uploads/images'
os.makedirs(app.config['IMAGE_UPLOAD_FOLDER'], exist_ok=True)

app.config['WORD_UPLOAD_FOLDER'] = 'uploads/word_documents'
os.makedirs(app.config['WORD_UPLOAD_FOLDER'], exist_ok=True)

# Email Configuration
mail = Mail(app)

def allowed_file(filename, allowed_extensions=ALLOWED_EXTENSIONS):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def sanitize(value):
    return value if isinstance(value, str) else str(value)

def clean_text(text):
    return unicodedata.normalize("NFKD", text)

def parse_birth_date(date_str):
    """생년월일 문자열을 date 객체로 변환"""
    if not date_str:
        return None
    
    # 여러 형식 지원
    formats = [
        '%Y/%m/%d',  # 2025/07/08
        '%Y-%m-%d',  # 2025-07-08
        '%Y.%m.%d',  # 2025.07.08
    ]
    
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt).date()
        except ValueError:
            continue
    
    # 모든 형식이 실패하면 None 반환
    return None

def parse_event_date(date_str):
    """이벤트 날짜 문자열을 date 객체로 변환 (YYYY/MM/DD 또는 YYYY-MM-DD 지원)"""
    for fmt in ('%Y/%m/%d', '%Y-%m-%d'):
        try:
            return datetime.strptime(date_str, fmt).date()
        except ValueError:
            continue
    raise ValueError(f"날짜 형식이 올바르지 않습니다: {date_str}")

def get_form_data(request):
    """폼 데이터를 수집하여 딕셔너리로 반환"""
    return {
        'name_kor': request.form.get('name_kor') or request.args.get('name_kor', ''),
        'email': request.form.get('email') or request.args.get('email', ''),
        'terms_service': request.form.get('terms_service') or request.args.get('terms_service', ''),
        'terms_privacy': request.form.get('terms_privacy') or request.args.get('terms_privacy', ''),
        'username': request.form.get('username', ''),
        'name_eng_first': request.form.get('name_eng_first', ''),
        'name_eng_last': request.form.get('name_eng_last', ''),
        'birth_date': request.form.get('birth_date', ''),
        'gender': request.form.get('gender', ''),
        'phone_prefix': request.form.get('phone_prefix', ''),
        'phone_middle': request.form.get('phone_middle', ''),
        'phone_end': request.form.get('phone_end', ''),
        'mobile_prefix': request.form.get('mobile_prefix', ''),
        'mobile_middle': request.form.get('mobile_middle', ''),
        'mobile_end': request.form.get('mobile_end', ''),
        'license_number': request.form.get('license_number', ''),
        'workplace_name': request.form.get('workplace_name', ''),
        'workplace_name_eng': request.form.get('workplace_name_eng', ''),
        'workplace_type': request.form.get('workplace_type', ''),
        'position': request.form.get('position', ''),
        'specialty': request.form.get('specialty', ''),
        'specialty_eng': request.form.get('specialty_eng', ''),
        'address': request.form.get('address', ''),
        'address_eng': request.form.get('address_eng', ''),
        'workplace_phone_prefix': request.form.get('workplace_phone_prefix', ''),
        'workplace_phone_middle': request.form.get('workplace_phone_middle', ''),
        'workplace_phone_end': request.form.get('workplace_phone_end', ''),
        'workplace_fax_prefix': request.form.get('workplace_fax_prefix', ''),
        'workplace_fax_middle': request.form.get('workplace_fax_middle', ''),
        'workplace_fax_end': request.form.get('workplace_fax_end', ''),
        'mail_receipt_location': request.form.get('mail_receipt_location', ''),
        'home_address': request.form.get('home_address', ''),
        'alma_mater': request.form.get('alma_mater', ''),
        'major': request.form.get('major', ''),
        'graduation_year': request.form.get('graduation_year', ''),
        'highest_degree': request.form.get('highest_degree', ''),
        'residency_institution': request.form.get('residency_institution', ''),
        'residency_hospital': request.form.get('residency_hospital', ''),
        'specialist_year': request.form.get('specialist_year', ''),
        'profile_public': request.form.get('profile_public', ''),
        'sms_receipt': request.form.get('sms_receipt', ''),
        'email_receipt': request.form.get('email_receipt', ''),
        'mail_receipt': request.form.get('mail_receipt', '')
    }

def init_db():
    """데이터베이스 초기화"""
    with app.app_context():
        db.create_all()
        print("Database tables created!")

def get_all_participants_with_attendance():
    """모든 참가자 정보 조회 (출석 정보 포함)"""
    participants = Participant.query.all()
    return [{
        'id': p.id, 'event_id': p.event_id, 'registration': p.registration or '', 
        'division': p.division or '', 'role': p.role or '', 'country': p.country or '', 
        'name_kor': p.name_kor or '', 'affiliation_kor': p.affiliation_kor or '', 
        'department_kor': p.department_kor or '', 'first_name': p.first_name or '', 
        'family_name': p.family_name or '', 'affiliation_eng': p.affiliation_eng or '', 
        'department_eng': p.department_eng or '', 'accept_or_decline': p.accept_or_decline or '', 
        'email': p.email or '', 'phone': p.phone or '', 'position': p.position or '', 
        'license_number': str(p.license_number) if p.license_number else '', 
        'cv': p.cv or '', 'photo': p.photo or '', 'ppt': p.ppt or '', 'script': p.script or '', 
        'agree': p.agree or '', 'remark_user': str(p.remark_user) if p.remark_user else '', 
        'remark_admin': str(p.remark_admin) if p.remark_admin else '', 
        'check_in_time': p.check_in_time.isoformat() if p.check_in_time else '', 
        'check_out_time': p.check_out_time.isoformat() if p.check_out_time else '', 
        'decline_reason': p.decline_reason or ''
    } for p in participants]

def migrate_existing_events():
    """기존 이벤트에 event_id 생성 및 location 컬럼 추가, members 테이블에 first_name, last_name 컬럼 추가"""
    try:
        # 먼저 location 컬럼이 있는지 확인하고 없으면 추가
        try:
            with db.engine.connect() as connection:
                connection.execute(db.text("ALTER TABLE events ADD COLUMN location VARCHAR(500);"))
                connection.commit()
            print("Added location column to events table")
        except Exception as e:
            if "already exists" in str(e) or "duplicate column" in str(e).lower():
                print("Location column already exists")
            else:
                print(f"Error adding location column: {e}")
        
        # members 테이블에 first_name, last_name, department_kor, department_eng 컬럼 추가
        try:
            with db.engine.connect() as connection:
                connection.execute(db.text("ALTER TABLE members ADD COLUMN first_name VARCHAR(100);"))
                connection.execute(db.text("ALTER TABLE members ADD COLUMN last_name VARCHAR(100);"))
                connection.execute(db.text("ALTER TABLE members ADD COLUMN department_kor VARCHAR(200);"))
                connection.execute(db.text("ALTER TABLE members ADD COLUMN department_eng VARCHAR(200);"))
                connection.commit()
            print("Added first_name, last_name, department_kor, department_eng columns to members table")
        except Exception as e:
            if "already exists" in str(e) or "duplicate column" in str(e).lower():
                print("members table columns already exist")
            else:
                print(f"Error adding members table columns: {e}")
        
        # participants 테이블에 name_eng 컬럼 추가
        try:
            with db.engine.connect() as connection:
                connection.execute(db.text("ALTER TABLE participants ADD COLUMN name_eng VARCHAR(100);"))
                connection.commit()
            print("Added name_eng column to participants table")
        except Exception as e:
            if "already exists" in str(e) or "duplicate column" in str(e).lower():
                print("name_eng column already exists in participants table")
            else:
                print(f"Error adding name_eng column to participants table: {e}")
        
        # participants 테이블에 country_code, birth_date, workplace_type 컬럼 추가
        try:
            with db.engine.connect() as connection:
                connection.execute(db.text("ALTER TABLE participants ADD COLUMN country_code VARCHAR(10);"))
                connection.execute(db.text("ALTER TABLE participants ADD COLUMN birth_date DATE;"))
                connection.execute(db.text("ALTER TABLE participants ADD COLUMN workplace_type VARCHAR(50);"))
                connection.commit()
            print("Added country_code, birth_date, workplace_type columns to participants table")
        except Exception as e:
            if "already exists" in str(e) or "duplicate column" in str(e).lower():
                print("participants table new columns already exist")
            else:
                print(f"Error adding participants table new columns: {e}")
        
        # 기존 이벤트에 event_id가 없는 경우 자동 생성
        events = Event.query.filter(Event.event_id.is_(None)).all()
        for event in events:
            event_date = event.start_date
            base_event_id = event_date.strftime('%Y-%m')
            existing_events = Event.query.filter(
                Event.event_id.like(f"{base_event_id}%"),
                Event.id != event.id
            ).all()
            
            if not existing_events:
                new_event_id = base_event_id
            else:
                suffixes = [e.event_id[len(base_event_id) + 1:] for e in existing_events if e.event_id.startswith(base_event_id + '-')]
                if not suffixes:
                    new_event_id = f"{base_event_id}-A"
                else:
                    last_suffix = sorted(suffixes)[-1]
                    next_suffix = chr(ord(last_suffix) + 1)
                    new_event_id = f"{base_event_id}-{next_suffix}"
            
            event.event_id = new_event_id
        
        db.session.commit()
        print(f"Migrated {len(events)} events with new event_id")
    except Exception as e:
        print(f"Migration error: {e}")
        db.session.rollback()

def migrate_participant_codes():
    """참가자 코드 생성"""
    events = Event.query.all()
    for event in events:
        participants = Participant.query.filter_by(event_id=event.id).order_by(Participant.id).all()
        for code, participant in enumerate(participants, 1):
            participant.code = code
    
    db.session.commit()

# 데이터베이스 초기화
with app.app_context():
    init_db()
    migrate_existing_events()
    migrate_participant_codes()

@app.route('/')
def index():
    """루트 경로 - 이벤트 페이지로 리다이렉트"""
    return redirect(url_for('admin_event_page'))

@app.route('/compose_email')
def compose_email():
    participant_ids = request.args.get('participants', '')
    return render_template('compose_email.html', participant_ids=participant_ids)

@app.route('/event')
def admin_event_page():
    """행사 관리 페이지"""
    events = db.session.query(
        Event.id, Event.name, Event.start_date, Event.end_date, Event.event_id,
        db.func.count(Participant.id).label('participant_count')
    ).outerjoin(Participant).group_by(
        Event.id, Event.name, Event.start_date, Event.end_date, Event.event_id
    ).order_by(Event.start_date.asc()).all()
    
    return render_template('admin_event.html', events=events)

@app.route('/members')
def members_page():
    """회원 목록 페이지"""
    members = Member.query.order_by(Member.created_at.desc()).all()
    response = make_response(render_template('admin_members.html', members=members))
    # 캐시 방지 헤더 추가
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

@app.route('/add_member', methods=['GET', 'POST'])
def add_member():
    """관리자용 회원 추가 페이지"""
    if request.method == 'POST':
        # 폼 데이터 수집
        form_data = get_form_data(request)
        
        try:
            # 필수 필드 검증
            required_fields = ['username', 'password', 'license_number']
            for field in required_fields:
                if not request.form.get(field):
                    return render_template('add_member.html', 
                                         error=f'{field}을(를) 입력해주세요.',
                                         **form_data)
            
            # 중복 확인
            existing_member = Member.query.filter(
                (Member.username == request.form.get('username')) | 
                (Member.email == request.form.get('email'))
            ).first()
            
            if existing_member:
                return render_template('add_member.html', 
                                     error='이미 등록된 아이디 또는 이메일입니다.',
                                     **form_data)
            
            # 새 회원 생성
            member = Member(
                username=request.form.get('username'),
                password_hash=generate_password_hash(request.form.get('password')),
                email=request.form.get('email'),
                name_kor=request.form.get('name_kor'),
                name_eng=f"{request.form.get('name_eng_first', '')} {request.form.get('name_eng_last', '')}".strip(),
                first_name=request.form.get('name_eng_first', ''),
                last_name=request.form.get('name_eng_last', ''),
                birth_date=parse_birth_date(request.form.get('birth_date')) if request.form.get('birth_date') else None,
                gender=request.form.get('gender'),
                phone=f"{request.form.get('phone_prefix', '')}-{request.form.get('phone_middle', '')}-{request.form.get('phone_end', '')}" if request.form.get('phone_prefix') else None,
                mobile=f"{request.form.get('mobile_prefix', '')}-{request.form.get('mobile_middle', '')}-{request.form.get('mobile_end', '')}" if request.form.get('mobile_prefix') else None,
                license_number=request.form.get('license_number'),
                workplace_name=request.form.get('workplace_name'),
                workplace_name_eng=request.form.get('workplace_name_eng'),
                workplace_type=request.form.get('workplace_type'),
                position=request.form.get('position'),
                specialty=request.form.get('specialty'),
                specialty_eng=request.form.get('specialty_eng'),
                address=request.form.get('address'),
                address_eng=request.form.get('address_eng'),
                workplace_phone=f"{request.form.get('workplace_phone_prefix')}-{request.form.get('workplace_phone_middle')}-{request.form.get('workplace_phone_end')}" if request.form.get('workplace_phone_prefix') else None,
                workplace_fax=f"{request.form.get('workplace_fax_prefix')}-{request.form.get('workplace_fax_middle')}-{request.form.get('workplace_fax_end')}" if request.form.get('workplace_fax_prefix') else None,
                mail_receipt_location=request.form.get('mail_receipt_location'),
                home_address=request.form.get('home_address'),
                alma_mater=request.form.get('alma_mater'),
                major=request.form.get('major'),
                graduation_year=int(request.form.get('graduation_year')) if request.form.get('graduation_year') else None,
                highest_degree=request.form.get('highest_degree'),
                residency_institution=request.form.get('residency_institution'),
                residency_hospital=request.form.get('residency_hospital'),
                specialist_year=int(request.form.get('specialist_year')) if request.form.get('specialist_year') else None,
                profile_public=request.form.get('profile_public') == 'true',
                sms_receipt=request.form.get('sms_receipt') == 'true',
                email_receipt=request.form.get('email_receipt') == 'true',
                mail_receipt=request.form.get('mail_receipt') == 'true',
                terms_agreed=True,
                terms_agreed_at=datetime.utcnow(),
                is_active=True  # 관리자가 추가하는 회원은 바로 활성화
            )
            
            # 프로필 사진 업로드
            if 'profile_photo' in request.files:
                file = request.files['profile_photo']
                if file and file.filename and allowed_file(file.filename, ALLOWED_IMAGE_EXTENSIONS):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config['IMAGE_UPLOAD_FOLDER'], f"profile_{member.username}_{filename}")
                    file.save(filepath)
                    member.profile_photo = filepath
            
            db.session.add(member)
            db.session.commit()
            
            return redirect(url_for('members_page'))
            
        except Exception as e:
            db.session.rollback()
            logging.error(f"Error adding member: {str(e)}")
            return render_template('add_member.html', 
                                 error=f'회원 추가 중 오류가 발생했습니다: {str(e)}',
                                 **form_data)
    
    # GET 요청 시 빈 폼 표시
    return render_template('add_member.html')

@app.route('/register/step1', methods=['GET', 'POST'])
def register_step1():
    """회원가입 1단계: 가입확인"""
    if request.method == 'POST':
        name_kor = request.form.get('name_kor')
        email = request.form.get('email')
        
        if not name_kor or not email:
            return render_template('register_step1.html', error='이름과 이메일을 모두 입력해주세요.')
        
        # 기존 회원 재확인
        member = Member.query.filter(
            (Member.name_kor == name_kor) & (Member.email == email)
        ).first()
        
        if member:
            return render_template('register_step1.html', error='이미 등록된 회원입니다.')
        
        # 2단계로 이동
        return redirect(url_for('register_step2', name_kor=name_kor, email=email))
    
    # GET 요청 시 URL 파라미터에서 값 가져오기
    name_kor = request.args.get('name_kor', '')
    email = request.args.get('email', '')
    
    return render_template('register_step1.html', name_kor=name_kor, email=email)

@app.route('/register/step2', methods=['GET', 'POST'])
def register_step2():
    """회원가입 2단계: 약관동의"""
    if request.method == 'POST':
        name_kor = request.form.get('name_kor')
        email = request.form.get('email')
        terms_service = request.form.get('terms_service')
        terms_privacy = request.form.get('terms_privacy')
        
        if not terms_service or not terms_privacy:
            return render_template('register_step2.html', 
                                 error='모든 약관에 동의해주세요.',
                                 name_kor=name_kor, email=email)
        
        # 3단계로 이동
        return redirect(url_for('register_step3', 
                               name_kor=name_kor, 
                               email=email,
                               terms_service=terms_service,
                               terms_privacy=terms_privacy))
    
    # GET 요청 시 URL 파라미터에서 값 가져오기
    name_kor = request.args.get('name_kor', '')
    email = request.args.get('email', '')
    
    return render_template('register_step2.html', name_kor=name_kor, email=email)

@app.route('/register/step3', methods=['GET', 'POST'])
def register_step3():
    """회원가입 3단계: 회원정보 입력"""
    if request.method == 'POST':
        # 폼 데이터 수집
        form_data = get_form_data(request)
        
        try:
            # 필수 필드 검증
            required_fields = ['username', 'password', 'password_confirm', 'license_number']
            for field in required_fields:
                if not request.form.get(field):
                    return render_template('register_step3.html', 
                                         error=f'{field}을(를) 입력해주세요.',
                                         **form_data)
            
            # 비밀번호 확인
            if request.form.get('password') != request.form.get('password_confirm'):
                return render_template('register_step3.html', 
                                     error='비밀번호가 일치하지 않습니다.',
                                     **form_data)
            
            # 중복 확인
            existing_member = Member.query.filter(
                (Member.username == request.form.get('username')) | 
                (Member.email == request.form.get('email'))
            ).first()
            
            if existing_member:
                return render_template('register_step3.html', 
                                     error='이미 등록된 아이디 또는 이메일입니다.',
                                     **form_data)
            
            # 새 회원 생성
            member = Member(
                username=request.form.get('username'),
                password_hash=generate_password_hash(request.form.get('password')),
                email=request.form.get('email'),
                name_kor=request.form.get('name_kor'),
                name_eng=f"{request.form.get('name_eng_first', '')} {request.form.get('name_eng_last', '')}".strip(),
                first_name=request.form.get('name_eng_first', ''),
                last_name=request.form.get('name_eng_last', ''),
                birth_date=parse_birth_date(request.form.get('birth_date')) if request.form.get('birth_date') else None,
                gender=request.form.get('gender'),
                phone=f"{request.form.get('phone_prefix', '')}-{request.form.get('phone_middle', '')}-{request.form.get('phone_end', '')}" if request.form.get('phone_prefix') else None,
                mobile=f"{request.form.get('mobile_prefix', '')}-{request.form.get('mobile_middle', '')}-{request.form.get('mobile_end', '')}" if request.form.get('mobile_prefix') else None,
                license_number=request.form.get('license_number'),
                workplace_name=request.form.get('workplace_name'),
                workplace_name_eng=request.form.get('workplace_name_eng'),
                workplace_type=request.form.get('workplace_type'),
                position=request.form.get('position'),
                specialty=request.form.get('specialty'),
                specialty_eng=request.form.get('specialty_eng'),
                address=request.form.get('address'),
                address_eng=request.form.get('address_eng'),
                workplace_phone=f"{request.form.get('workplace_phone_prefix')}-{request.form.get('workplace_phone_middle')}-{request.form.get('workplace_phone_end')}" if request.form.get('workplace_phone_prefix') else None,
                workplace_fax=f"{request.form.get('workplace_fax_prefix')}-{request.form.get('workplace_fax_middle')}-{request.form.get('workplace_fax_end')}" if request.form.get('workplace_fax_prefix') else None,
                mail_receipt_location=request.form.get('mail_receipt_location'),
                home_address=request.form.get('home_address'),
                alma_mater=request.form.get('alma_mater'),
                major=request.form.get('major'),
                graduation_year=int(request.form.get('graduation_year')) if request.form.get('graduation_year') else None,
                highest_degree=request.form.get('highest_degree'),
                residency_institution=request.form.get('residency_institution'),
                residency_hospital=request.form.get('residency_hospital'),
                specialist_year=int(request.form.get('specialist_year')) if request.form.get('specialist_year') else None,
                profile_public=request.form.get('profile_public') == 'true',
                sms_receipt=request.form.get('sms_receipt') == 'true',
                email_receipt=request.form.get('email_receipt') == 'true',
                mail_receipt=request.form.get('mail_receipt') == 'true',
                terms_agreed=True,
                terms_agreed_at=datetime.utcnow()
            )
            
            # 프로필 사진 업로드
            if 'profile_photo' in request.files:
                file = request.files['profile_photo']
                if file and file.filename and allowed_file(file.filename, ALLOWED_IMAGE_EXTENSIONS):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config['IMAGE_UPLOAD_FOLDER'], f"profile_{member.username}_{filename}")
                    file.save(filepath)
                    member.profile_photo = filepath
            
            db.session.add(member)
            db.session.commit()
            
            # 4단계로 이동
            return redirect(url_for('register_step4', success='회원가입이 완료되었습니다. 관리자 승인 후 로그인이 가능합니다.'))
            
        except Exception as e:
            db.session.rollback()
            logging.error(f"Error during registration: {str(e)}")
            return render_template('register_step3.html', 
                                 error=f'회원가입 중 오류가 발생했습니다: {str(e)}',
                                 **form_data)
    
    # GET 요청 시 URL 파라미터에서 값 가져오기
    form_data = get_form_data(request)
    return render_template('register_step3.html', **form_data)

@app.route('/register/step4')
def register_step4():
    """회원가입 4단계: 가입완료"""
    success = request.args.get('success', '')
    return render_template('register_step4.html', success=success)

@app.route('/api/check-username', methods=['POST'])
def check_username():
    """아이디 중복확인 API"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        
        if not username:
            return jsonify({'available': False, 'error': '아이디를 입력해주세요.'})
        
        # 아이디 형식 검사 (4~20자 영문소문자, 숫자)
        import re
        if not re.match(r'^[a-z0-9]{4,20}$', username):
            return jsonify({'available': False, 'error': '아이디는 4~20자의 영문소문자와 숫자만 사용 가능합니다.'})
        
        # 데이터베이스에서 중복 확인
        existing_member = Member.query.filter_by(username=username).first()
        
        if existing_member:
            return jsonify({'available': False, 'message': '이미 사용 중인 아이디입니다.'})
        else:
            return jsonify({'available': True, 'message': '사용 가능한 아이디입니다.'})
            
    except Exception as e:
        logging.error(f"Error checking username: {str(e)}")
        return jsonify({'available': False, 'error': '중복확인 중 오류가 발생했습니다.'})

@app.route('/api/check-license', methods=['POST'])
def check_license():
    """의사면허번호 중복확인 API"""
    try:
        data = request.get_json()
        license_number = data.get('license_number', '').strip()
        
        if not license_number:
            return jsonify({'available': False, 'error': '의사면허번호를 입력해주세요.'})
        
        # 데이터베이스에서 중복 확인
        existing_member = Member.query.filter_by(license_number=license_number).first()
        
        if existing_member:
            return jsonify({'available': False, 'message': '이미 등록된 의사면허번호입니다.'})
        else:
            return jsonify({'available': True, 'message': '사용 가능한 의사면허번호입니다.'})
            
    except Exception as e:
        logging.error(f"Error checking license: {str(e)}")
        return jsonify({'available': False, 'error': '중복확인 중 오류가 발생했습니다.'})

@app.route('/members/delete', methods=['POST'])
def delete_members():
    """회원 삭제"""
    selected_members = request.form.getlist('selected_members')
    if selected_members:
        member_ids = [int(member_id) for member_id in selected_members]
        Member.query.filter(Member.id.in_(member_ids)).delete(synchronize_session=False)
        db.session.commit()
    return redirect(url_for('members_page'))

@app.route('/members/edit/<int:member_id>', methods=['GET', 'POST'])
def edit_member(member_id):
    """회원 정보 수정"""
    member = Member.query.get_or_404(member_id)
    
    if request.method == 'POST':
        try:
            # 아이디 중복 검사 (현재 아이디와 다를 때만)
            new_username = request.form.get('username')
            if new_username != member.username:
                existing_member = Member.query.filter(Member.username == new_username).first()
                if existing_member:
                    return render_template('edit_member.html', member=member, error='이미 사용 중인 아이디입니다.')
            
            # 폼 데이터로 회원 정보 업데이트
            member.username = new_username
            member.email = request.form.get('email')
            member.name_kor = request.form.get('name_kor')
            member.name_eng = f"{request.form.get('name_eng_first', '')} {request.form.get('name_eng_last', '')}".strip()
            member.first_name = request.form.get('name_eng_first', '')
            member.last_name = request.form.get('name_eng_last', '')
            member.birth_date = parse_birth_date(request.form.get('birth_date')) if request.form.get('birth_date') else None
            member.gender = request.form.get('gender')
            member.phone = f"{request.form.get('phone_prefix', '')}-{request.form.get('phone_middle', '')}-{request.form.get('phone_end', '')}" if request.form.get('phone_prefix') else None
            member.mobile = f"{request.form.get('mobile_prefix', '')}-{request.form.get('mobile_middle', '')}-{request.form.get('mobile_end', '')}" if request.form.get('mobile_prefix') else None
            member.license_number = request.form.get('license_number')
            member.workplace_name = request.form.get('workplace_name')
            member.workplace_name_eng = request.form.get('workplace_name_eng')
            member.workplace_type = request.form.get('workplace_type')
            member.position = request.form.get('position')
            member.specialty = request.form.get('specialty')
            member.specialty_eng = request.form.get('specialty_eng')
            member.address = request.form.get('address')
            member.address_eng = request.form.get('address_eng')
            member.workplace_phone = f"{request.form.get('workplace_phone_prefix')}-{request.form.get('workplace_phone_middle')}-{request.form.get('workplace_phone_end')}" if request.form.get('workplace_phone_prefix') else None
            member.workplace_fax = f"{request.form.get('workplace_fax_prefix')}-{request.form.get('workplace_fax_middle')}-{request.form.get('workplace_fax_end')}" if request.form.get('workplace_fax_prefix') else None
            member.mail_receipt_location = request.form.get('mail_receipt_location')
            member.home_address = request.form.get('home_address')
            member.alma_mater = request.form.get('alma_mater')
            member.major = request.form.get('major')
            member.graduation_year = int(request.form.get('graduation_year')) if request.form.get('graduation_year') else None
            member.highest_degree = request.form.get('highest_degree')
            member.residency_institution = request.form.get('residency_institution')
            member.residency_hospital = request.form.get('residency_hospital')
            member.specialist_year = int(request.form.get('specialist_year')) if request.form.get('specialist_year') else None
            member.profile_public = request.form.get('profile_public') == 'true'
            member.sms_receipt = request.form.get('sms_receipt') == 'true'
            member.email_receipt = request.form.get('email_receipt') == 'true'
            member.mail_receipt = request.form.get('mail_receipt') == 'true'
            
            # 비밀번호 변경 (입력된 경우에만)
            if request.form.get('password'):
                if request.form.get('password') != request.form.get('password_confirm'):
                    return render_template('edit_member.html', member=member, error='비밀번호가 일치하지 않습니다.')
                member.password_hash = generate_password_hash(request.form.get('password'))
            
            # 프로필 사진 업로드
            if 'profile_photo' in request.files:
                file = request.files['profile_photo']
                if file and allowed_file(file.filename, ALLOWED_IMAGE_EXTENSIONS):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config['IMAGE_UPLOAD_FOLDER'], f"profile_{member.username}_{filename}")
                    file.save(filepath)
                    member.profile_photo = filepath
            
            db.session.commit()
            return redirect(url_for('members_page'))
            
        except Exception as e:
            db.session.rollback()
            logging.error(f"Error updating member: {str(e)}")
            return render_template('edit_member.html', member=member, error=f'회원 정보 수정 중 오류가 발생했습니다: {str(e)}')
    
    # GET 요청 시 회원 정보를 폼에 미리 채움
    return render_template('edit_member.html', member=member)

@app.route('/select_columns_members')
def select_columns_members():
    """회원 Excel 내보내기 컬럼 선택 페이지"""
    members = request.args.get('members', '')
    return render_template('select_columns_members.html', members=members)

@app.route('/download_members_excel', methods=['POST'])
def download_members_excel():
    """선택된 회원을 선택된 컬럼으로 Excel 다운로드"""
    try:
        # 선택된 회원 ID
        member_ids = request.form.get('selected_members', '').split(',')
        member_ids = [int(mid) for mid in member_ids if mid]
        
        # 선택된 컬럼
        selected_columns = request.form.getlist('selected_columns')
        
        if not member_ids:
            return jsonify({'error': 'No members selected'}), 400
        
        if not selected_columns:
            return jsonify({'error': 'No columns selected'}), 400
        
        # 회원 조회
        members = Member.query.filter(Member.id.in_(member_ids)).all()
        
        # 표준 컬럼 매핑
        column_mapping = {
            '아이디': lambda m: m.username,
            '성별': lambda m: m.gender,
            '성명(KOR)': lambda m: m.name_kor,
            '성명(ENG)': lambda m: m.name_eng,
            '이름(First Name)': lambda m: m.first_name,
            '성(Last Name)': lambda m: m.last_name,
            '이메일': lambda m: m.email,
            '전화': lambda m: m.phone,
            '소속(ENG)': lambda m: m.workplace_name_eng,
            '과(ENG)': lambda m: m.specialty_eng,
            '소속(KOR)': lambda m: m.workplace_name,
            '과(KOR)': lambda m: m.specialty,
            '직위': lambda m: m.position,
            '면허번호': lambda m: m.license_number,
            '생년월일': lambda m: m.birth_date.strftime('%Y-%m-%d') if m.birth_date else '',
            '회원구분': lambda m: m.workplace_type
        }
        
        # DataFrame 생성
        data = []
        for member in members:
            row = {}
            for col in selected_columns:
                if col in column_mapping:
                    row[col] = column_mapping[col](member) or ''
            data.append(row)
        
        df = pd.DataFrame(data)
        
        # 선택된 컬럼 순서대로 재정렬
        df = df[selected_columns]
        
        # Excel 파일 생성
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='회원목록', index=False)
        
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'회원목록_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        )
    
    except Exception as e:
        logging.error(f"Error in download_members_excel: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/upload_members', methods=['GET', 'POST'])
def upload_members():
    """회원 일괄 업로드"""
    if request.method == 'GET':
        return render_template('upload_members.html')
    
    if request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected'}), 400
        
        if file and allowed_file(file.filename, {'xlsx', 'xls', 'csv'}):
            try:
                # 파일 읽기
                if file.filename.endswith('.csv'):
                    df = pd.read_csv(file, encoding='utf-8', dtype=str)
                else:
                    df = pd.read_excel(file, dtype=str)
                
                # NaN 값을 빈 문자열로 변환
                df = df.fillna('')
                
                # 컬럼명 정규화 (공백 제거, 대소문자 통일)
                df.columns = [col.strip() for col in df.columns]
                print(f"파일 업로드: 총 {len(df)}행, {len(df.columns)}개 컬럼")
                
                success_count = 0
                failure_count = 0
                errors = []
                
                # 컬럼 매핑 정의 (Excel 컬럼명 -> 데이터베이스 필드명)
                # 이미지의 실제 컬럼명들을 우선적으로 포함
                column_mapping = {
                    # 아이디 관련 (이미지에는 없지만 시스템에서 사용)
                    '아이디': 'username',
                    'id': 'username',
                    'user_id': 'username',
                    'username': 'username',
                    
                    # 성명 관련 (이미지 컬럼명 포함)
                    '성명(KOR)': 'name_kor',
                    '성명': 'name_kor',
                    'name_kor': 'name_kor',
                    '한글성명': 'name_kor',
                    '한국어성명': 'name_kor',
                    
                    '성명(ENG)': 'name_eng',
                    'name_eng': 'name_eng',
                    '영문성명': 'name_eng',
                    'english_name': 'name_eng',
                    
                    # First/Last Name (이미지 컬럼명)
                    '이름(First Name)': 'first_name',
                    'First Name': 'first_name',
                    'first_name': 'first_name',
                    'firstname': 'first_name',
                    '이름': 'first_name',
                    
                    '성(Last Name)': 'last_name',
                    'Last Name': 'last_name',
                    'last_name': 'last_name',
                    'lastname': 'last_name',
                    '성': 'last_name',
                    'family_name': 'last_name',
                    
                    # 이메일 관련 (이미지 컬럼명)
                    '이메일': 'email',
                    'Email': 'email',
                    'email': 'email',
                    'e-mail': 'email',
                    'mail': 'email',
                    
                    # 전화 관련 (이미지 컬럼명)
                    '전화': 'phone',
                    'Phone': 'phone',
                    'phone': 'phone',
                    '핸드폰': 'phone',
                    'mobile': 'phone',
                    '휴대전화': 'phone',
                    '전화번호': 'phone',
                    'office_phone': 'phone',
                    
                    # 소속 관련 (이미지 컬럼명)
                    '소속(KOR)': 'workplace_name',
                    '소속(ENG)': 'workplace_name_eng',
                    '근무처': 'workplace_name',
                    'workplace_name': 'workplace_name',
                    'company': 'workplace_name',
                    '회사': 'workplace_name',
                    '근무처(영문)': 'workplace_name_eng',
                    'workplace_name_eng': 'workplace_name_eng',
                    'company_eng': 'workplace_name_eng',
                    
                    # 과/부서 관련 (이미지 컬럼명)
                    '과(KOR)': 'specialty',
                    '과(ENG)': 'specialty_eng',
                    '과': 'specialty',
                    'specialty': 'specialty',
                    '전문분야': 'specialty',
                    '과(영문)': 'specialty_eng',
                    'specialty_eng': 'specialty_eng',
                    '전문분야_영문': 'specialty_eng',
                    '부서': 'specialty',
                    'department_kor': 'specialty',
                    '부서명': 'specialty',
                    '부서(영문)': 'specialty_eng',
                    'department_eng': 'specialty_eng',
                    '부서명_영문': 'specialty_eng',
                    
                    # 기타 필드들
                    '성별': 'gender',
                    'gender': 'gender',
                    'sex': 'gender',
                    
                    '직위': 'position',
                    'Position': 'position',
                    'position': 'position',
                    'job_title': 'position',
                    
                    '면허번호': 'license_number',
                    'License Number': 'license_number',
                    'license_number': 'license_number',
                    'license': 'license_number',
                    
                    '생년월일': 'birth_date',
                    'Birth Date': 'birth_date',
                    'birth_date': 'birth_date',
                    'birth': 'birth_date',
                    '생일': 'birth_date',
                    
                    '회원구분': 'workplace_type',
                    'Membership Type': 'workplace_type',
                    'workplace_type': 'workplace_type',
                    'member_type': 'workplace_type',
                    
                    # 추가 전화 필드들
                    '근무처전화': 'workplace_phone',
                    'workplace_phone': 'workplace_phone',
                    'company_phone': 'workplace_phone'
                }
                
                # 각 행 처리
                for index, row in df.iterrows():
                    try:
                        # 필수 필드 검증 (유연한 컬럼명 매칭)
                        username = ''
                        name_kor = ''
                        email = ''
                        
                        # 아이디 찾기 (이미지에는 없지만 시스템에서 사용)
                        for col in row.index:
                            col_lower = col.strip().lower()
                            if col_lower in ['아이디', 'id', 'user_id', 'username']:
                                username = str(row[col]).strip()
                                break
                        
                        # 성명(KOR) 찾기
                        for col in row.index:
                            col_lower = col.strip().lower()
                            if col_lower in ['성명(kor)', '성명', 'name_kor', '한글성명', '한국어성명']:
                                name_kor = str(row[col]).strip()
                                if name_kor and name_kor.lower() not in ['nan', 'none', 'null', '']:
                                    break
                                else:
                                    name_kor = ''  # 빈 값이면 다시 빈 문자열로 설정
                        
                        # 이메일 찾기
                        for col in row.index:
                            col_lower = col.strip().lower()
                            if col_lower in ['email', '이메일', 'e-mail', 'mail']:
                                email = str(row[col]).strip()
                                if email and email.lower() not in ['nan', 'none', 'null', '']:
                                    break
                                else:
                                    email = ''  # 빈 값이면 다시 빈 문자열로 설정
                        
                        # 필수 필드 검증 (빈 문자열도 무효로 처리)
                        if (not username or username == '') and (not name_kor or name_kor == '') and (not email or email == ''):
                            errors.append(f"행 {index + 2}: 최소 하나의 식별 정보(아이디, 성명(KOR), Email)가 필요합니다.")
                            failure_count += 1
                            continue
                        
                        # 중복 검사 (아이디, 이메일, 면허번호 확인)
                        existing_member = None
                        if username:
                            existing_member = Member.query.filter(Member.username == username).first()
                        if not existing_member and email:
                            existing_member = Member.query.filter(Member.email == email).first()
                        
                        # 면허번호 중복 검사 (유연한 컬럼명 매칭)
                        license_number = ''
                        for col in row.index:
                            col_lower = col.strip().lower()
                            if col_lower in ['면허번호', 'license number', 'license_number', 'license']:
                                license_number = str(row[col]).strip()
                                break
                        
                        if license_number and not existing_member:
                            existing_member = Member.query.filter(Member.license_number == license_number).first()
                        
                        if existing_member:
                            if username and existing_member.username == username:
                                errors.append(f"행 {index + 2}: 아이디 '{username}'가 이미 존재합니다.")
                            elif email and existing_member.email == email:
                                errors.append(f"행 {index + 2}: Email '{email}'가 이미 존재합니다.")
                            elif license_number and existing_member.license_number == license_number:
                                errors.append(f"행 {index + 2}: 면허번호 '{license_number}'가 이미 존재합니다.")
                            else:
                                errors.append(f"행 {index + 2}: 중복된 정보가 있습니다.")
                            failure_count += 1
                            continue
                        
                        # 데이터 매핑 (Excel에 있는 컬럼만)
                        member_data = {
                            'password_hash': generate_password_hash('temp123!'),  # 임시 비밀번호 해시
                            'is_active': True,
                            'terms_agreed': True  # 관리자가 업로드했으므로 동의한 것으로 간주
                        }
                        
                        # 디버깅: 첫 3행만 출력
                        if index < 3:
                            print(f"[샘플] 행 {index + 2} 처리 중...")
                        
                        for excel_col, db_field in column_mapping.items():
                            matched_col = None
                            
                            # 1. 정확한 매칭 시도
                            if excel_col in row:
                                matched_col = excel_col
                            else:
                                # 2. 유연한 매칭 시도 (대소문자, 공백, 특수문자 무시)
                                excel_col_normalized = excel_col.strip().lower().replace(' ', '').replace('(', '').replace(')', '')
                                for actual_col in row.index:
                                    actual_col_normalized = actual_col.strip().lower().replace(' ', '').replace('(', '').replace(')', '')
                                    if actual_col_normalized == excel_col_normalized:
                                        matched_col = actual_col
                                        break
                            
                            if matched_col:
                                value = row[matched_col]
                                if pd.isna(value) or value is None:
                                    value = ''
                                else:
                                    value = str(value).strip()
                                
                                # 빈 문자열이나 NaN 값이 아닌 경우에만 저장
                                if value and value.lower() not in ['nan', 'none', 'null', '']:
                                    if db_field == 'birth_date':
                                        member_data[db_field] = parse_date(value)
                                    else:
                                        member_data[db_field] = value
                                    # 첫 3행만 매핑 정보 출력
                                    if index < 3:
                                        print(f"  ✅ 매핑됨: '{matched_col}' -> '{db_field}' = '{value[:50]}'")  # 값은 최대 50자만
                        
                        # 첫 3행만 최종 데이터 출력
                        if index < 3:
                            print(f"최종 member_data 키: {list(member_data.keys())}")
                        
                        # 필수 필드가 없으면 기본값 설정
                        if not member_data.get('username'):
                            # 이메일이 있으면 이메일 앞부분을 아이디로 사용, 없으면 자동 생성
                            if member_data.get('email'):
                                email_prefix = member_data['email'].split('@')[0]
                                member_data['username'] = f"{email_prefix}_{datetime.now().strftime('%m%d%H%M')}"
                            else:
                                member_data['username'] = f"user_{index + 1}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                        
                        # name_kor가 없으면 기본값 설정 (필수 필드)
                        if not member_data.get('name_kor'):
                            # name_eng이 있으면 그것을 사용, 없으면 first_name + last_name 조합 사용
                            if member_data.get('name_eng'):
                                member_data['name_kor'] = member_data['name_eng']
                            elif member_data.get('first_name') and member_data.get('last_name'):
                                member_data['name_kor'] = f"{member_data['last_name']} {member_data['first_name']}"
                            else:
                                member_data['name_kor'] = f"회원_{index + 1}"
                        
                        # 이메일이 없으면 임시 이메일 생성
                        if not member_data.get('email'):
                            member_data['email'] = f"temp{index + 1}@example.com"
                        
                        # 새 회원 생성
                        member = Member(**member_data)
                        
                        db.session.add(member)
                        success_count += 1
                        
                    except Exception as e:
                        errors.append(f"행 {index + 2}: {str(e)}")
                        failure_count += 1
                        # 개별 행 오류 시 세션 롤백
                        db.session.rollback()
                        continue
                
                # 데이터베이스 저장
                try:
                    db.session.commit()
                    print(f"✅ 업로드 완료: 성공 {success_count}건, 실패 {failure_count}건")
                except Exception as e:
                    db.session.rollback()
                    print(f"❌ 데이터베이스 저장 실패: {str(e)}")
                    return jsonify({'success': False, 'message': f'데이터베이스 저장 중 오류: {str(e)}'}), 500
                
                return jsonify({
                    'success': True,
                    'success_count': success_count,
                    'failure_count': failure_count,
                    'errors': errors[:10]  # 최대 10개 에러만 반환
                })
                
            except Exception as e:
                db.session.rollback()
                return jsonify({'success': False, 'message': f'파일 처리 중 오류가 발생했습니다: {str(e)}'}), 500
        
        return jsonify({'success': False, 'message': '지원하지 않는 파일 형식입니다.'}), 400

def parse_date(date_str):
    """날짜 문자열을 파싱"""
    if not date_str or date_str == '':
        return None
    
    date_str = str(date_str).strip()
    
    # Excel 날짜 숫자 형식 처리 (예: 42567)
    try:
        if date_str.isdigit():
            # Excel 날짜는 1900년 1월 1일부터의 일수
            excel_date = int(date_str)
            if 1 <= excel_date <= 73050:  # 1900-2099 범위
                base_date = datetime(1900, 1, 1)
                return (base_date + timedelta(days=excel_date - 2)).date()  # Excel은 1900년을 윤년으로 잘못 계산
    except (ValueError, OverflowError):
        pass
    
    # 일반적인 날짜 형식들
    date_formats = [
        '%Y-%m-%d',      # 1986-03-12
        '%Y/%m/%d',      # 1986/03/12
        '%m/%d/%Y',      # 03/12/1986
        '%d/%m/%Y',      # 12/03/1986
        '%Y.%m.%d',      # 1986.03.12
        '%m-%d-%Y',      # 03-12-1986
        '%d-%m-%Y',      # 12-03-1986
        '%Y년%m월%d일',   # 1986년03월12일
        '%Y%m%d',        # 19860312
    ]
    
    for fmt in date_formats:
        try:
            return datetime.strptime(date_str, fmt).date()
        except ValueError:
            continue
    
    # pandas의 dateutil.parser 사용 (마지막 수단)
    try:
        from dateutil import parser
        parsed_date = parser.parse(date_str)
        return parsed_date.date()
    except:
        # 디버그 모드에서만 로그 출력
        # print(f"날짜 파싱 실패: '{date_str}'")
        return None

@app.route('/add_member_as_participant/<int:event_id>', methods=['POST'])
def add_member_as_participant(event_id):
    """회원을 참가자로 등록"""
    try:
        data = request.get_json()
        member_ids = data.get('member_ids', [])
        
        if not member_ids:
            return jsonify({'success': False, 'message': '선택된 회원이 없습니다.'}), 400
        
        event = Event.query.get_or_404(event_id)
        added_count = 0
        skipped_count = 0
        
        for member_id in member_ids:
            member = Member.query.get(member_id)
            if not member:
                continue
            
            # 이미 해당 이벤트에 참가자로 등록되어 있는지 확인
            existing_participant = Participant.query.filter_by(
                event_id=event_id,
                email=member.email
            ).first()
            
            if existing_participant:
                skipped_count += 1
                continue
            
            # 참가자 코드 생성 (현재 이벤트의 최대 코드 + 1)
            max_code = db.session.query(db.func.max(Participant.code)).filter_by(event_id=event_id).scalar()
            next_code = (max_code or 0) + 1
            
            # 성명(ENG) 자동 조합
            name_eng_value = member.name_eng
            if not name_eng_value:
                if member.first_name and member.last_name:
                    name_eng_value = f"{member.first_name} {member.last_name}"
                elif member.first_name:
                    name_eng_value = member.first_name
                elif member.last_name:
                    name_eng_value = member.last_name
            
            # 새 참가자 생성
            participant = Participant(
                event_id=event_id,
                code=next_code,
                name_kor=member.name_kor,
                name_eng=name_eng_value,
                first_name=member.first_name,
                family_name=member.last_name,
                affiliation_kor=member.workplace_name,
                affiliation_eng=member.workplace_name_eng,
                department_kor=member.department_kor,
                department_eng=member.department_eng,
                email=member.email,
                phone=member.mobile or member.phone,
                position=member.position,
                license_number=member.license_number,
                birth_date=member.birth_date,
                workplace_type=member.workplace_type,
                accept_or_decline='대기'
            )
            
            db.session.add(participant)
            added_count += 1
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'{added_count}명의 회원이 참가자로 등록되었습니다.',
            'added_count': added_count,
            'skipped_count': skipped_count
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/members', methods=['GET'])
def get_members_api():
    """회원 목록 API (참가자 등록용)"""
    try:
        search = request.args.get('search', '')
        members = Member.query
        
        if search:
            members = members.filter(
                (Member.name_kor.contains(search)) |
                (Member.name_eng.contains(search)) |
                (Member.email.contains(search)) |
                (Member.workplace_name.contains(search))
            )
        
        members = members.limit(50).all()  # 최대 50명만 반환
        
        members_data = []
        for member in members:
            members_data.append({
                'id': member.id,
                'name_kor': member.name_kor,
                'name_eng': member.name_eng,
                'first_name': member.first_name,
                'last_name': member.last_name,
                'email': member.email,
                'workplace_name': member.workplace_name,
                'workplace_name_eng': member.workplace_name_eng,
                'department_kor': member.department_kor,
                'department_eng': member.department_eng,
                'position': member.position,
                'mobile': member.mobile,
                'license_number': member.license_number
            })
        
        return jsonify({'success': True, 'members': members_data})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'오류가 발생했습니다: {str(e)}'}), 500

@app.route('/create_event', methods=['POST'])
def create_event():
    """이벤트 생성"""
    name = request.form['name']
    start_date = parse_event_date(request.form['start_date'])
    end_date = parse_event_date(request.form['end_date'])
    
    # event_id 생성
    base_event_id = start_date.strftime('%Y-%m')
    existing_events = Event.query.filter(Event.event_id.like(f"{base_event_id}%")).all()
    
    if not existing_events:
        event_id = base_event_id
    else:
        suffixes = [e.event_id[len(base_event_id) + 1:] for e in existing_events if e.event_id.startswith(base_event_id + '-')]
        if not suffixes:
            event_id = f"{base_event_id}-A"
        else:
            last_suffix = sorted(suffixes)[-1]
            next_suffix = chr(ord(last_suffix) + 1)
            event_id = f"{base_event_id}-{next_suffix}"
    
    new_event = Event(name=name, start_date=start_date, end_date=end_date, event_id=event_id)
    db.session.add(new_event)
    db.session.commit()
    
    return redirect(url_for('admin_event_page'))

@app.route('/get_files/<int:participant_id>', methods=['GET'])
def get_files(participant_id):
    """참가자 파일 목록 조회"""
    files = ParticipantFile.query.filter_by(participant_id=participant_id).all()
    return jsonify([{
        'id': f.id,
        'filename': f.filename,
        'filepath': f.filepath,
        'file_type': f.file_type,
        'file_size': f.file_size
    } for f in files])

@app.route('/download_file/<int:file_id>', methods=['GET'])
def download_file(file_id):
    """파일 다운로드"""
    file_record = ParticipantFile.query.get_or_404(file_id)
    try:
        return send_file(file_record.filepath, as_attachment=True, download_name=file_record.filename)
    except Exception as e:
        logging.error(f"File download failed: {e}")
        return "File not found", 404

@app.route('/delete_events', methods=['POST'])
def delete_events():
    """이벤트 삭제"""
    selected_events = request.form.getlist('selected_events')
    if selected_events:
        # 문자열을 정수로 변환
        event_ids = [int(event_id) for event_id in selected_events]
        Event.query.filter(Event.id.in_(event_ids)).delete(synchronize_session=False)
        db.session.commit()
    return redirect(url_for('admin_event_page'))

@app.route('/event/<int:event_id>')
def participant_management(event_id):
    """참가자 관리 페이지"""
    event = Event.query.get_or_404(event_id)
    participants = Participant.query.filter_by(event_id=event_id).all()
    
    # 튜플 형태로 변환 (기존 템플릿 호환성)
    event_tuple = (event.id, event.name, event.start_date, event.end_date, event.event_id, event.location)
    participants_tuples = []
    
    for p in participants:
        participants_tuples.append((
            p.id, p.event_id, p.code, p.role, p.country, p.country_code,
            p.name_kor, p.name_eng, p.first_name, p.family_name, p.email, p.phone,
            p.affiliation_eng, p.department_eng, p.affiliation_kor, p.department_kor,
            p.position, p.license_number, p.birth_date, p.workplace_type, p.registration,
            p.accept_or_decline, p.cv, p.photo, p.ppt, p.script, p.agree,
            p.remark_user, p.remark_admin, p.check_in_time, p.check_out_time, p.decline_reason
        ))
    
    return render_template('admin_participants.html', event=event_tuple, participants=participants_tuples)

@app.route('/add_participant/<int:event_id>', methods=['GET', 'POST'])
def add_participant(event_id):
    """참가자 추가"""
    event = Event.query.get_or_404(event_id)
    
    if request.method == 'POST':
        try:
            # 생년월일 파싱
            birth_date_value = None
            if request.form.get('birth_date'):
                birth_date_value = parse_date(request.form.get('birth_date'))
            
            # 성명(ENG) 자동 조합
            name_eng_value = request.form.get('name_eng')
            if not name_eng_value:
                first_name = request.form.get('first_name')
                family_name = request.form.get('family_name')
                if first_name and family_name:
                    name_eng_value = f"{first_name} {family_name}"
                elif first_name:
                    name_eng_value = first_name
                elif family_name:
                    name_eng_value = family_name
            
            # 폼 데이터 처리
            participant = Participant(
                event_id=event_id,
                role=request.form.get('role'),
                country=request.form.get('country'),
                country_code=request.form.get('country_code'),
                name_kor=request.form.get('name_kor'),
                name_eng=name_eng_value,
                first_name=request.form.get('first_name'),
                family_name=request.form.get('family_name'),
                email=request.form.get('email'),
                phone=request.form.get('phone'),
                affiliation_eng=request.form.get('affiliation_eng'),
                department_eng=request.form.get('department_eng'),
                affiliation_kor=request.form.get('affiliation_kor'),
                department_kor=request.form.get('department_kor'),
                position=request.form.get('position'),
                license_number=request.form.get('license_number'),
                birth_date=birth_date_value,
                workplace_type=request.form.get('workplace_type'),
                registration=request.form.get('registration'),
                accept_or_decline=request.form.get('accept_or_decline') or '대기',
                agree=request.form.get('agree'),
                remark_user=request.form.get('remark_user'),
                remark_admin=request.form.get('remark_admin')
            )
            
            # 파일 업로드 처리
            for field in ['cv', 'photo', 'ppt', 'script']:
                if field in request.files:
                    file = request.files[field]
                    if file and file.filename and allowed_file(file.filename):
                        filename = secure_filename(file.filename)
                        filepath = os.path.join(app.config['UPLOAD_FOLDER'], SUBFOLDERS.get(field, ''), filename)
                        file.save(filepath)
                        setattr(participant, field, filepath)
            
            db.session.add(participant)
            db.session.commit()
            
            # 참가자 코드 자동 생성
            if not participant.code:
                # 해당 이벤트의 최대 코드 번호 찾기
                max_code = db.session.query(db.func.max(Participant.code)).filter_by(event_id=event_id).scalar()
                participant.code = (max_code or 0) + 1
                db.session.commit()
            
            return redirect(url_for('participant_management', event_id=event_id))
            
        except Exception as e:
            db.session.rollback()
            logging.error(f"Error adding participant: {str(e)}")
            return render_template('add_participant.html', event=event, error=f"참가자 추가 중 오류가 발생했습니다: {str(e)}")
    
    return render_template('add_participant.html', event=event)

@app.route('/upload_participants/<int:event_id>', methods=['GET', 'POST'])
def upload_participants(event_id):
    """참가자 일괄 업로드"""
    try:
        logging.info(f"=== Upload participants called for event_id: {event_id} ===")
        logging.info(f"Request method: {request.method}")
        logging.info(f"Request content type: {request.content_type}")
        logging.info(f"Request content length: {request.content_length}")
        
        event = Event.query.get_or_404(event_id)
        logging.info(f"Event found: {event.name} (ID: {event.id}, event_id: {event.event_id})")
        
        def parse_datetime_or_none(val):
            try:
                if pd.isna(val) or val in ['', 'nan', 'NaN', None]:
                    return None
                return parser.parse(str(val))
            except Exception:
                return None
        
        if request.method == 'POST':
            logging.info(f"POST request received for event_id: {event_id}")
            logging.info(f"Request files keys: {list(request.files.keys())}")
            logging.info(f"Request form keys: {list(request.form.keys())}")
            
            if 'file' not in request.files:
                logging.error("No file uploaded")
                return jsonify({'error': 'No file uploaded'}), 400
            
            file = request.files['file']
            if file.filename == '':
                logging.error("No file selected")
                return jsonify({'error': 'No file selected'}), 400
            
            logging.info(f"Processing file: {file.filename} for event_id: {event_id}")
            
            if file and allowed_file(file.filename, {'csv', 'xlsx'}):
                try:
                    # 파일 내용을 메모리에 읽기
                    file_content = file.read()
                    file_size = len(file_content)
                    logging.info(f"File size: {file_size} bytes")
                    
                    if file_size > 100 * 1024 * 1024:  # 100MB
                        logging.error(f"File too large: {file_size} bytes")
                        return jsonify({'error': 'File too large. Maximum size is 100MB.'}), 400
                    
                    # 파일 내용을 BytesIO로 변환하여 pandas가 읽을 수 있도록 함
                    from io import BytesIO
                    file_stream = BytesIO(file_content)
                    if file.filename.endswith('.csv'):
                        df = pd.read_csv(file_stream, encoding='utf-8', dtype=str, na_values=['nan', 'NaN', ''], keep_default_na=False)
                    else:
                        df = pd.read_excel(file_stream, dtype=str, na_values=['nan', 'NaN', ''], keep_default_na=False)
                    
                    logging.info(f"File loaded successfully. Shape: {df.shape}")
                    logging.info(f"Original columns: {list(df.columns)}")
                    
                    # NaN 값을 빈 문자열로 변환
                    df = df.fillna('')
                    
                    # 컬럼명 정규화 (소문자로 변환하고 공백 제거)
                    df.columns = [col.lower().strip() for col in df.columns]
                    
                    # 이메일과 전화번호 정규화 (벡터화된 버전)
                    # 이메일 정규화: 소문자 변환 및 공백 제거
                    if 'email' in df.columns:
                        df['email_normalized'] = df['email'].fillna('').astype(str).str.strip().str.lower()
                        df['email_normalized'] = df['email_normalized'].replace(['nan', 'NaN', 'None'], '')
                    else:
                        df['email_normalized'] = ''
                    
                    # 전화번호 정규화: 하이픈, 공백, 괄호 제거
                    if 'phone' in df.columns:
                        df['phone_normalized'] = df['phone'].fillna('').astype(str).str.strip()
                        df['phone_normalized'] = df['phone_normalized'].str.replace('-', '', regex=False)
                        df['phone_normalized'] = df['phone_normalized'].str.replace(' ', '', regex=False)
                        df['phone_normalized'] = df['phone_normalized'].str.replace('(', '', regex=False)
                        df['phone_normalized'] = df['phone_normalized'].str.replace(')', '', regex=False)
                        df['phone_normalized'] = df['phone_normalized'].replace(['nan', 'NaN', 'None'], '')
                    else:
                        df['phone_normalized'] = ''
                    
                    logging.info(f"Normalized columns: {list(df.columns)}")
                    
                    # 컬럼명 매핑 (한글 컬럼명 지원)
                    column_mapping = {
                        '이벤트 id': 'event_id', '코드': 'code', '역할': 'role', '국가': 'country', '국가약어': 'country_code',
                        '성명(kor)': 'name_kor', '성명(eng)': 'name_eng', '이름(first name)': 'first_name', '성(last name)': 'family_name',
                        '이메일': 'email', '전화': 'phone', '소속(eng)': 'affiliation_eng', '과(eng)': 'department_eng',
                        '소속(kor)': 'affiliation_kor', '과(kor)': 'department_kor', '직위': 'position', '면허번호': 'license_number',
                        '생년월일': 'birth_date', '회원구분': 'workplace_type', '등록구분': 'registration', '승인/거절': 'accept_or_decline',
                        'cv': 'cv', '사진': 'photo', 'ppt': 'ppt', 'script': 'script', '동의여부': 'agree',
                        '비고(사용자)': 'remark_user', '비고(관리자)': 'remark_admin', '체크인': 'check_in_time', '체크아웃': 'check_out_time',
                        '거절 사유': 'decline_reason', '초청년도': 'invitation_year',
                        # 영문 컬럼명도 지원
                        'role': 'role', 'country': 'country', 'country code': 'country_code',
                        'name (kor)': 'name_kor', 'name (eng)': 'name_eng', 'first name': 'first_name', 'family name': 'family_name',
                        'email': 'email', 'phone': 'phone', 'affiliation (eng)': 'affiliation_eng', 'department (eng)': 'department_eng',
                        'affiliation (kor)': 'affiliation_kor', 'department (kor)': 'department_kor', 'position': 'position', 'license #': 'license_number',
                        'birth date': 'birth_date', 'workplace type': 'workplace_type', 'registration': 'registration', 'accept/decline': 'accept_or_decline',
                        'photo': 'photo', 'agree': 'agree', 'remark (user)': 'remark_user', 'remark (admin)': 'remark_admin',
                        'check-in': 'check_in_time', 'check-out': 'check_out_time', 'decline reason': 'decline_reason',
                        'invitation year': 'invitation_year'
                    }
                    df.rename(columns=column_mapping, inplace=True)
                    logging.info(f"After mapping columns: {list(df.columns)}")
                    
                    # 이메일과 전화번호 정규화는 이미 위에서 완료됨 (중복 제거)
                
                # 모든 가능한 컬럼
                    all_columns = [
                    'event_id', 'code', 'role', 'country', 'country_code', 'name_kor', 'name_eng', 'first_name', 'family_name',
                    'email', 'phone', 'affiliation_eng', 'department_eng', 'affiliation_kor', 'department_kor',
                    'position', 'license_number', 'birth_date', 'workplace_type', 'registration', 'accept_or_decline',
                    'cv', 'photo', 'ppt', 'script', 'agree', 'remark_user', 'remark_admin',
                    'check_in_time', 'check_out_time', 'decline_reason', 'invitation_year'
                    ]
                    present_columns = [col for col in all_columns if col in df.columns]
                    missing_columns = [col for col in all_columns if col not in df.columns]
                    
                    logging.info(f"Present columns: {present_columns}")
                    logging.info(f"Missing columns: {missing_columns}")
                    logging.info(f"First row data: {df.iloc[0].to_dict() if len(df) > 0 else 'No data'}")
                    
                    # 중복 처리: 다양한 기준으로 중복 감지 및 최신 초청년도 선택
                    # 초청년도 컬럼 확인 - 명시적 컬럼 또는 첫 번째 컬럼
                    year_column = None
                    if 'invitation_year' in df.columns:
                        year_column = 'invitation_year'
                        logging.info("초청년도 컬럼 발견 - 중복 처리 시작")
                    elif len(df.columns) > 0:
                        # 첫 번째 컬럼을 년도로 사용 (년도로 추정되는 경우만)
                        first_col = df.columns[0]
                        first_col_sample = df[first_col].dropna().head(5)
                        
                        # 첫 번째 컬럼이 년도인지 확인 (1900-2100 사이의 숫자)
                        is_year_column = True
                        for val in first_col_sample:
                            try:
                                year_val = int(str(val).strip())
                                if year_val < 1900 or year_val > 2100:
                                    is_year_column = False
                                    break
                            except:
                                is_year_column = False
                                break
                        
                        if is_year_column and len(first_col_sample) > 0:
                            year_column = first_col
                            logging.info(f"첫 번째 컬럼 '{first_col}'을 년도 컬럼으로 사용 - 중복 처리 시작")
                    
                    if year_column:
                        # 초청년도를 숫자로 변환하는 함수
                        def safe_int_conversion(val):
                            try:
                                if pd.isna(val) or val == '':
                                    return 0
                                return int(str(val).strip())
                            except:
                                return 0
                        
                        # 초청년도 정수 변환
                        df['invitation_year_int'] = df[year_column].apply(safe_int_conversion)
                        
                        latest_records = []
                        processed_indices = set()
                        
                        # 이름+이메일 조합을 위한 통합 이름 필드 생성
                        def get_unified_name(row):
                            """이름(KOR 또는 ENG)을 통합하여 반환"""
                            name_kor = str(row.get('name_kor', '')).strip() if pd.notna(row.get('name_kor')) else ''
                            name_eng = str(row.get('name_eng', '')).strip() if pd.notna(row.get('name_eng')) else ''
                            first_name = str(row.get('first_name', '')).strip() if pd.notna(row.get('first_name')) else ''
                            family_name = str(row.get('family_name', '')).strip() if pd.notna(row.get('family_name')) else ''
                            
                            # 우선순위: name_kor > name_eng > first_name + family_name
                            if name_kor:
                                return name_kor.lower()
                            elif name_eng:
                                return name_eng.lower()
                            elif first_name and family_name:
                                return f"{first_name} {family_name}".lower()
                            elif first_name:
                                return first_name.lower()
                            elif family_name:
                                return family_name.lower()
                            return ''
                        
                        df['unified_name'] = df.apply(get_unified_name, axis=1)
                        
                        # 1단계: 이메일이 있는 경우 - 이메일 + 이름 조합으로 중복 처리
                        email_df = df[(df['email_normalized'].notna()) & (df['email_normalized'] != '') & 
                                     (df['unified_name'] != '')]
                        
                        if len(email_df) > 0:
                            logging.info(f"이메일+이름 중복 처리: {len(email_df)}개 레코드")
                            # 이메일 + 통합 이름으로 그룹화
                            email_name_groups = email_df.groupby(['email_normalized', 'unified_name'])
                            
                            for (email, unified_name), group in email_name_groups:
                                if len(group) > 1:
                                    logging.info(f"중복 발견: {unified_name}({email}) - {len(group)}개 레코드")
                                    latest_record = group.loc[group['invitation_year_int'].idxmax()]
                                    latest_records.append(latest_record.to_dict())
                                    logging.info(f"  → {latest_record['invitation_year_int']}년 데이터 선택")
                                else:
                                    latest_records.append(group.iloc[0].to_dict())
                                
                                processed_indices.update(group.index.tolist())
                        
                        # 2단계: 이메일만으로 중복 처리 (이름이 없는 경우)
                        email_only_df = df[(df['email_normalized'].notna()) & (df['email_normalized'] != '') &
                                           (~df.index.isin(processed_indices))]
                        
                        if len(email_only_df) > 0:
                            logging.info(f"이메일만으로 중복 처리: {len(email_only_df)}개 레코드")
                            email_only_groups = email_only_df.groupby(['email_normalized'])
                            
                            for email, group in email_only_groups:
                                if len(group) > 1:
                                    logging.info(f"이메일 중복 발견: {email} - {len(group)}개 레코드")
                                    latest_record = group.loc[group['invitation_year_int'].idxmax()]
                                    latest_records.append(latest_record.to_dict())
                                    logging.info(f"  → {latest_record['invitation_year_int']}년 데이터 선택")
                                else:
                                    latest_records.append(group.iloc[0].to_dict())
                                
                                processed_indices.update(group.index.tolist())
                        
                        # 3단계: 성명(KOR) + 면허번호 기준 중복 처리 (이메일 없는 국내 참가자)
                        kor_license_df = df[(df['name_kor'].notna()) & (df['name_kor'] != '') & 
                                           (df['license_number'].notna()) & (df['license_number'] != '') &
                                           (~df.index.isin(processed_indices))]
                        
                        if len(kor_license_df) > 0:
                            logging.info(f"국내 참가자 중복 처리: {len(kor_license_df)}개 레코드")
                            kor_license_groups = kor_license_df.groupby(['name_kor', 'license_number'])
                            
                            for (name_kor, license_num), group in kor_license_groups:
                                if len(group) > 1:
                                    logging.info(f"국내 중복 발견: {name_kor}({license_num}) - {len(group)}개 레코드")
                                    latest_record = group.loc[group['invitation_year_int'].idxmax()]
                                    latest_records.append(latest_record.to_dict())
                                    logging.info(f"  → {latest_record['invitation_year_int']}년 데이터 선택")
                                else:
                                    latest_records.append(group.iloc[0].to_dict())
                                
                                processed_indices.update(group.index.tolist())
                        
                        # 4단계: 나머지 데이터 (중복 기준에 해당하지 않는 데이터)
                        remaining_df = df[~df.index.isin(processed_indices)]
                        if len(remaining_df) > 0:
                            logging.info(f"중복 기준에 해당하지 않는 데이터: {len(remaining_df)}개 레코드")
                            latest_records.extend(remaining_df.to_dict('records'))
                    
                        # 최종 DataFrame 재구성
                        if latest_records:
                            df = pd.DataFrame(latest_records)
                            # 임시 컬럼 제거
                            if 'invitation_year_int' in df.columns:
                                df = df.drop('invitation_year_int', axis=1)
                            if 'unified_name' in df.columns:
                                df = df.drop('unified_name', axis=1)
                            logging.info(f"중복 처리 완료: {len(df)}개 레코드로 축소")
                        else:
                            logging.warning("중복 처리 후 유효한 레코드가 없음")
                    else:
                        # 초청년도 컬럼이 없어도 이름+이메일+전화번호로 중복 처리
                        logging.info("초청년도 컬럼 없음 - 이름+이메일+전화번호로 중복 처리")
                        
                        latest_records = []
                        processed_indices = set()
                        
                        # 1단계: 성명(KOR) + 정규화된 이메일 + 정규화된 전화번호
                        if 'name_kor' in df.columns and 'email_normalized' in df.columns and 'phone_normalized' in df.columns:
                            kor_email_phone_df = df[(df['name_kor'].notna()) & (df['name_kor'] != '') & 
                                                   (df['email_normalized'].notna()) & (df['email_normalized'] != '') &
                                                   (df['phone_normalized'].notna()) & (df['phone_normalized'] != '')]
                            
                            if len(kor_email_phone_df) > 0:
                                logging.info(f"이름+이메일+전화번호 중복 처리: {len(kor_email_phone_df)}개 레코드")
                                kor_email_phone_groups = kor_email_phone_df.groupby(['name_kor', 'email_normalized', 'phone_normalized'])
                                
                                for (name_kor, email, phone), group in kor_email_phone_groups:
                                    if len(group) > 1:
                                        logging.info(f"중복 발견: {name_kor}({email}, {phone}) - {len(group)}개 레코드")
                                        # 첫 번째 레코드 선택
                                        latest_records.append(group.iloc[0].to_dict())
                                    else:
                                        latest_records.append(group.iloc[0].to_dict())
                                    processed_indices.update(group.index.tolist())
                        
                        # 2단계: 성명(KOR) + 정규화된 이메일 (전화번호 없는 경우)
                        if 'name_kor' in df.columns and 'email_normalized' in df.columns:
                            kor_email_df = df[(df['name_kor'].notna()) & (df['name_kor'] != '') & 
                                            (df['email_normalized'].notna()) & (df['email_normalized'] != '') &
                                            (~df.index.isin(processed_indices))]
                            
                            if len(kor_email_df) > 0:
                                logging.info(f"이름+이메일 중복 처리: {len(kor_email_df)}개 레코드")
                                kor_email_groups = kor_email_df.groupby(['name_kor', 'email_normalized'])
                                
                                for (name_kor, email), group in kor_email_groups:
                                    if len(group) > 1:
                                        logging.info(f"중복 발견: {name_kor}({email}) - {len(group)}개 레코드")
                                        latest_records.append(group.iloc[0].to_dict())
                                    else:
                                        latest_records.append(group.iloc[0].to_dict())
                                    processed_indices.update(group.index.tolist())
                        
                        # 3단계: 정규화된 이메일만으로 중복 처리
                        if 'email_normalized' in df.columns:
                            email_only_df = df[(df['email_normalized'].notna()) & (df['email_normalized'] != '') &
                                              (~df.index.isin(processed_indices))]
                            
                            if len(email_only_df) > 0:
                                logging.info(f"이메일만으로 중복 처리: {len(email_only_df)}개 레코드")
                                email_only_groups = email_only_df.groupby(['email_normalized'])
                                
                                for email, group in email_only_groups:
                                    if len(group) > 1:
                                        logging.info(f"이메일 중복 발견: {email} - {len(group)}개 레코드")
                                        latest_records.append(group.iloc[0].to_dict())
                                    else:
                                        latest_records.append(group.iloc[0].to_dict())
                                    processed_indices.update(group.index.tolist())
                        
                        # 4단계: 나머지 데이터
                        remaining_df = df[~df.index.isin(processed_indices)]
                        if len(remaining_df) > 0:
                            logging.info(f"중복 기준에 해당하지 않는 데이터: {len(remaining_df)}개 레코드")
                            latest_records.extend(remaining_df.to_dict('records'))
                        
                        # 최종 DataFrame 재구성
                        if latest_records:
                            df = pd.DataFrame(latest_records)
                            logging.info(f"중복 처리 완료: {len(df)}개 레코드로 축소")
                        else:
                            logging.warning("중복 처리 후 유효한 레코드가 없음")
                    
                    # 데이터 처리 및 저장
                    # 성능 최적화: 기존 참가자를 한 번에 로드하여 메모리에서 검색
                    logging.info("기존 참가자 데이터 로드 중...")
                    existing_participants_list = Participant.query.filter_by(event_id=event_id).all()
                    logging.info(f"기존 참가자 {len(existing_participants_list)}명 로드 완료")
                    
                    # 참가자 코드 자동 생성을 위한 최대 코드 번호 미리 계산
                    max_code = db.session.query(db.func.max(Participant.code)).filter_by(event_id=event_id).scalar()
                    next_code = (max_code or 0) + 1
                    logging.info(f"다음 참가자 코드 시작 번호: {next_code}")
                    
                    # 검색을 위한 인덱스 생성 (메모리에서 빠른 검색)
                    participants_by_name_email = {}  # key: (name_kor or name_eng, email_lower)
                    participants_by_email = {}  # key: email_lower
                    participants_by_name_phone = {}  # key: (name_kor, phone_normalized)
                    
                    for p in existing_participants_list:
                        email_lower = (p.email or '').strip().lower()
                        name_kor = (p.name_kor or '').strip()
                        name_eng = (p.name_eng or '').strip()
                        phone_normalized = (p.phone or '').replace('-', '').replace(' ', '').replace('(', '').replace(')', '')
                        
                        # 한글 이름 + 이메일 인덱스
                        if name_kor and email_lower:
                            participants_by_name_email[(name_kor, email_lower)] = p
                        # 영문 이름 + 이메일 인덱스
                        if name_eng and email_lower:
                            participants_by_name_email[(name_eng, email_lower)] = p
                        # 이메일만 인덱스
                        if email_lower:
                            if email_lower not in participants_by_email:
                                participants_by_email[email_lower] = p
                        # 이름 + 전화번호 인덱스
                        if name_kor and phone_normalized:
                            participants_by_name_phone[(name_kor, phone_normalized)] = p
                    
                    logging.info("인덱스 생성 완료, 참가자 처리 시작...")
                    
                    participants_added = 0
                    participants_updated = 0
                    code_counter = next_code  # 코드 카운터
                    
                    # CSV 내에서 이미 처리한 참가자 추적 (중복 방지)
                    processed_csv_participants = set()  # (name_kor or name_eng, email) 튜플 저장
                    
                    # iterrows() 대신 to_dict('records') 사용 (더 빠름)
                    df_dict_list = df.to_dict('records')
                    for row_dict in df_dict_list:
                        try:
                            sanitized_data = {}
                            for key in all_columns:
                                value = row_dict.get(key, '')
                                if isinstance(value, str):
                                    value = value.strip()
                                elif pd.isna(value) or value == '':
                                    value = ''
                                else:
                                    value = str(value).strip()
                                # datetime 필드의 경우 빈 값이면 None으로 설정
                                if key in ['check_in_time', 'check_out_time']:
                                    sanitized_data[key] = value if value else None
                                else:
                                    sanitized_data[key] = value
                            
                            # 최소한의 필수 데이터 확인
                            if not sanitized_data.get('name_kor') and not sanitized_data.get('first_name'):
                                logging.warning(f"Row: Missing name data, skipping")
                                continue
                            
                            # 생년월일 파싱
                            birth_date_value = None
                            if sanitized_data.get('birth_date'):
                                birth_date_value = parse_date(sanitized_data.get('birth_date'))
                            
                            # code 필드 처리 (빈 문자열을 None으로, 없으면 자동 생성)
                            code_value = sanitized_data.get('code')
                            if code_value == '' or code_value is None:
                                code_value = code_counter
                                code_counter += 1
                            else:
                                try:
                                    code_value = int(code_value)
                                except (ValueError, TypeError):
                                    code_value = code_counter
                                    code_counter += 1
                            
                            # 성명(ENG) 자동 조합
                            name_eng_value = sanitized_data.get('name_eng', '')
                            first_name = sanitized_data.get('first_name', '')
                            family_name = sanitized_data.get('family_name', '')
                            
                            if not name_eng_value:
                                if first_name and family_name:
                                    name_eng_value = f"{first_name} {family_name}"
                                elif first_name:
                                    name_eng_value = first_name
                                elif family_name:
                                    name_eng_value = family_name
                            
                            # 초청년도 처리 (참고용으로만 사용, DB에는 저장하지 않음)
                            invitation_year_value = sanitized_data.get('invitation_year', '')
                            
                            # CSV 내에서 이미 처리한 참가자인지 먼저 확인
                            name_kor_value = sanitized_data.get('name_kor', '').strip()
                            email_value = sanitized_data.get('email', '').strip().lower() if sanitized_data.get('email') else ''
                            phone_value = sanitized_data.get('phone', '')
                            phone_normalized = phone_value.replace('-', '').replace(' ', '').replace('(', '').replace(')', '') if phone_value else ''
                            
                            # CSV 내 중복 체크: 이미 처리한 참가자인지 확인
                            csv_duplicate_key = None
                            if name_kor_value and email_value:
                                csv_duplicate_key = (name_kor_value, email_value)
                            elif name_eng_value and email_value:
                                csv_duplicate_key = (name_eng_value, email_value)
                            elif email_value:
                                csv_duplicate_key = (email_value, '')  # 이메일만으로도 체크
                            
                            if csv_duplicate_key and csv_duplicate_key in processed_csv_participants:
                                logging.info(f"CSV 내 중복 발견, 건너뜀: {name_kor_value or name_eng_value} ({email_value})")
                                continue  # 이미 처리한 참가자는 건너뛰기
                            
                            # 메모리에서 중복 체크: DB에 이미 존재하는 참가자인지 확인
                            existing_participant = None
                            
                            # 1순위: 한글 이름 + 이메일로 검색
                            if name_kor_value and email_value:
                                existing_participant = participants_by_name_email.get((name_kor_value, email_value))
                                if existing_participant:
                                    logging.debug(f"기존 참가자 발견 (한글이름+이메일): {name_kor_value} ({email_value})")
                            
                            # 2순위: 영문 이름 + 이메일로 검색
                            if not existing_participant and name_eng_value and email_value:
                                existing_participant = participants_by_name_email.get((name_eng_value, email_value))
                                if existing_participant:
                                    logging.debug(f"기존 참가자 발견 (영문이름+이메일): {name_eng_value} ({email_value})")
                            
                            # 3순위: 이메일만으로 검색 (이름이 다른 경우)
                            if not existing_participant and email_value:
                                existing_participant = participants_by_email.get(email_value)
                                if existing_participant:
                                    logging.debug(f"기존 참가자 발견 (이메일만): {email_value}")
                            
                            # 4순위: 이름 + 전화번호로 검색 (이메일이 없는 경우)
                            if not existing_participant and name_kor_value and phone_normalized:
                                existing_participant = participants_by_name_phone.get((name_kor_value, phone_normalized))
                                if existing_participant:
                                    logging.debug(f"기존 참가자 발견 (이름+전화번호): {name_kor_value} ({phone_value})")
                            
                            if existing_participant:
                                # 기존 참가자 업데이트
                                logging.info(f"기존 참가자 업데이트: ID {existing_participant.id}")
                                existing_participant.code = code_value
                                existing_participant.role = sanitized_data.get('role', '')
                                existing_participant.country = sanitized_data.get('country', '')
                                existing_participant.country_code = sanitized_data.get('country_code', '')
                                existing_participant.name_kor = name_kor_value
                                existing_participant.name_eng = name_eng_value
                                existing_participant.first_name = first_name
                                existing_participant.family_name = family_name
                                existing_participant.phone = sanitized_data.get('phone', '')
                                existing_participant.affiliation_eng = sanitized_data.get('affiliation_eng', '')
                                existing_participant.department_eng = sanitized_data.get('department_eng', '')
                                existing_participant.affiliation_kor = sanitized_data.get('affiliation_kor', '')
                                existing_participant.department_kor = sanitized_data.get('department_kor', '')
                                existing_participant.position = sanitized_data.get('position', '')
                                existing_participant.license_number = sanitized_data.get('license_number', '')
                                existing_participant.birth_date = birth_date_value
                                existing_participant.workplace_type = sanitized_data.get('workplace_type', '')
                                existing_participant.registration = sanitized_data.get('registration', '')
                                existing_participant.accept_or_decline = sanitized_data.get('accept_or_decline', '')
                                existing_participant.cv = sanitized_data.get('cv', '')
                                existing_participant.photo = sanitized_data.get('photo', '')
                                existing_participant.ppt = sanitized_data.get('ppt', '')
                                existing_participant.script = sanitized_data.get('script', '')
                                existing_participant.agree = sanitized_data.get('agree', '')
                                existing_participant.remark_user = sanitized_data.get('remark_user', '')
                                existing_participant.remark_admin = sanitized_data.get('remark_admin', '')
                                existing_participant.check_in_time = sanitized_data.get('check_in_time') if sanitized_data.get('check_in_time') else None
                                existing_participant.check_out_time = sanitized_data.get('check_out_time') if sanitized_data.get('check_out_time') else None
                                existing_participant.decline_reason = sanitized_data.get('decline_reason', '')
                                participants_updated += 1
                                
                                # CSV 내 처리 추적에 추가
                                if csv_duplicate_key:
                                    processed_csv_participants.add(csv_duplicate_key)
                                # 기존 참가자는 add 불필요 (이미 DB에 있음)
                            else:
                                # 새 참가자 생성
                                logging.info(f"새 참가자 추가: {name_kor_value or name_eng_value} ({email_value})")
                                participant = Participant(
                                    event_id=event_id,
                                    code=code_value,
                                    role=sanitized_data.get('role', ''),
                                    country=sanitized_data.get('country', ''),
                                    country_code=sanitized_data.get('country_code', ''),
                                    name_kor=name_kor_value,
                                    name_eng=name_eng_value,
                                    first_name=first_name,
                                    family_name=family_name,
                                    email=email_value,
                                    phone=sanitized_data.get('phone', ''),
                                    affiliation_eng=sanitized_data.get('affiliation_eng', ''),
                                    department_eng=sanitized_data.get('department_eng', ''),
                                    affiliation_kor=sanitized_data.get('affiliation_kor', ''),
                                    department_kor=sanitized_data.get('department_kor', ''),
                                    position=sanitized_data.get('position', ''),
                                    license_number=sanitized_data.get('license_number', ''),
                                    birth_date=birth_date_value,
                                    workplace_type=sanitized_data.get('workplace_type', ''),
                                    registration=sanitized_data.get('registration', ''),
                                    accept_or_decline=sanitized_data.get('accept_or_decline', ''),
                                    cv=sanitized_data.get('cv', ''),
                                    photo=sanitized_data.get('photo', ''),
                                    ppt=sanitized_data.get('ppt', ''),
                                    script=sanitized_data.get('script', ''),
                                    agree=sanitized_data.get('agree', ''),
                                    remark_user=sanitized_data.get('remark_user', ''),
                                    remark_admin=sanitized_data.get('remark_admin', ''),
                                    check_in_time=sanitized_data.get('check_in_time') if sanitized_data.get('check_in_time') else None,
                                    check_out_time=sanitized_data.get('check_out_time') if sanitized_data.get('check_out_time') else None,
                                    decline_reason=sanitized_data.get('decline_reason', '')
                                )
                                db.session.add(participant)
                                participants_added += 1
                                
                                # 로깅에 초청년도 정보 포함
                                name_display = sanitized_data.get('name_kor', sanitized_data.get('first_name', 'Unknown'))
                                year_info = f" (초청년도: {invitation_year_value})" if invitation_year_value else ""
                                logging.info(f"Added participant {participants_added}: {name_display}{year_info}")
                                
                                # CSV 내 처리 추적에 추가
                                if csv_duplicate_key:
                                    processed_csv_participants.add(csv_duplicate_key)
                        except Exception as e:
                            logging.error(f"Error processing row: {e}")
                            continue
                    
                    logging.info(f"Attempting to commit changes to database: {participants_added} added, {participants_updated} updated")
                    db.session.commit()
                    logging.info(f"Successfully committed: {participants_added} added, {participants_updated} updated")
                    
                    # 결과 메시지 생성
                    result_message = f'업로드 완료 - 추가: {participants_added}명, 업데이트: {participants_updated}명'
                    if missing_columns:
                        result_message += f' (누락 컬럼: {", ".join(missing_columns)})'
                        logging.warning(f"Missing columns: {missing_columns}")
                    
                    logging.info(f"Upload completed successfully. Added: {participants_added}, Updated: {participants_updated}")
                    
                    # JSON 응답 반환 (AJAX 요청인 경우)
                    return jsonify({
                        'success': True,
                        'message': result_message,
                        'added': participants_added,
                        'updated': participants_updated,
                        'missing_columns': missing_columns if missing_columns else []
                    }), 200
                except Exception as e:
                    import traceback
                    error_traceback = traceback.format_exc()
                    logging.error(f"Upload failed: {e}")
                    logging.error(f"Traceback: {error_traceback}")
                    db.session.rollback()
                    # 에러 응답을 JSON으로 반환하여 클라이언트가 제대로 받을 수 있도록 함
                    try:
                        return jsonify({'error': f'Upload failed: {str(e)}', 'details': error_traceback}), 500
                    except:
                        return f'Upload failed: {str(e)}', 500
            else:
                logging.error(f"Invalid file type: {file.filename}")
                return jsonify({'error': 'Invalid file type. Please upload CSV or Excel file.'}), 400
        else:
            # GET 요청 - 업로드 페이지 표시
            return render_template('upload_participant.html', event=event)
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logging.error(f"Upload participants function error: {e}")
        logging.error(f"Traceback: {error_traceback}")
        try:
            return jsonify({'error': f'Server error: {str(e)}'}), 500
        except:
            return f'Server error: {str(e)}', 500
        else:
            logging.error(f"Invalid file type: {file.filename}")
            return 'Invalid file type. Please upload CSV or Excel file.', 400
    
    return render_template('upload_participant.html', event=event)

@app.route('/delete_file/<int:file_id>', methods=['POST'])
def delete_file(file_id):
    """파일 삭제"""
    file_record = ParticipantFile.query.get_or_404(file_id)
    try:
        if os.path.exists(file_record.filepath):
            os.remove(file_record.filepath)
        db.session.delete(file_record)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        logging.error(f"File deletion failed: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/delete_participants/<int:event_id>', methods=['POST'])
def delete_participants(event_id):
    """참가자 삭제"""
    selected_participants = request.form.getlist('selected_participants')
    if selected_participants:
        # 문자열을 정수로 변환
        participant_ids = [int(participant_id) for participant_id in selected_participants]
        Participant.query.filter(Participant.id.in_(participant_ids)).delete(synchronize_session=False)
        db.session.commit()
    return redirect(url_for('participant_management', event_id=event_id))

@app.route('/edit_participant/<int:participant_id>', methods=['GET', 'POST'])
def edit_participant(participant_id):
    """참가자 편집"""
    participant = Participant.query.get_or_404(participant_id)
    
    if request.method == 'POST':
        # 생년월일 파싱
        birth_date_value = None
        if request.form.get('birth_date'):
            birth_date_value = parse_date(request.form.get('birth_date'))
        
        # 성명(ENG) 자동 조합
        name_eng_value = request.form.get('name_eng')
        first_name = request.form.get('first_name')
        family_name = request.form.get('family_name')
        
        if not name_eng_value:
            if first_name and family_name:
                name_eng_value = f"{first_name} {family_name}"
            elif first_name:
                name_eng_value = first_name
            elif family_name:
                name_eng_value = family_name
        
        # 폼 데이터 업데이트
        participant.role = request.form.get('role')
        participant.country = request.form.get('country')
        participant.country_code = request.form.get('country_code')
        participant.name_kor = request.form.get('name_kor')
        participant.name_eng = name_eng_value
        participant.first_name = first_name
        participant.family_name = family_name
        participant.email = request.form.get('email')
        participant.phone = request.form.get('phone')
        participant.affiliation_eng = request.form.get('affiliation_eng')
        participant.department_eng = request.form.get('department_eng')
        participant.affiliation_kor = request.form.get('affiliation_kor')
        participant.department_kor = request.form.get('department_kor')
        participant.position = request.form.get('position')
        participant.license_number = request.form.get('license_number')
        participant.birth_date = birth_date_value
        participant.workplace_type = request.form.get('workplace_type')
        participant.registration = request.form.get('registration')
        participant.accept_or_decline = request.form.get('accept_or_decline')
        participant.agree = request.form.get('agree')
        participant.remark_user = request.form.get('remark_user')
        participant.remark_admin = request.form.get('remark_admin')
        
        # 파일 업로드 처리
        for field in ['cv', 'photo', 'ppt', 'script']:
            if field in request.files:
                file = request.files[field]
                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], SUBFOLDERS.get(field, ''), filename)
                    file.save(filepath)
                    setattr(participant, field, filepath)
        
        db.session.commit()
        return redirect(url_for('participant_management', event_id=participant.event_id))
    
    return render_template('edit_participant.html', participant=participant)

@app.route('/delete_file_field/<int:participant_id>/<string:field>', methods=['POST'])
def delete_file_field(participant_id, field):
    """특정 파일 필드 삭제"""
    participant = Participant.query.get_or_404(participant_id)
    if hasattr(participant, field):
        filepath = getattr(participant, field)
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
        setattr(participant, field, None)
        db.session.commit()
    return jsonify({'success': True})

@app.route('/upload_image', methods=['POST'])
def upload_image():
    """이미지 업로드 (이메일용)"""
    try:
        if 'image' not in request.files:
            return jsonify({'status': 'error', 'message': '이미지가 업로드되지 않았습니다.'}), 400
        
        file = request.files['image']
        if file.filename == '':
            return jsonify({'status': 'error', 'message': '파일이 선택되지 않았습니다.'}), 400
        
        if file and allowed_file(file.filename, ALLOWED_IMAGE_EXTENSIONS):
            filename = secure_filename(file.filename)
            # 중복 파일명 방지를 위해 타임스탬프 추가
            import time
            timestamp = int(time.time() * 1000)
            name, ext = os.path.splitext(filename)
            filename = f"{name}_{timestamp}{ext}"
            filepath = os.path.join(app.config['IMAGE_UPLOAD_FOLDER'], filename)
            file.save(filepath)
            image_url = f'/uploads/images/{filename}'
            return jsonify({'status': 'success', 'image_url': image_url})
        
        return jsonify({'status': 'error', 'message': '지원하지 않는 파일 형식입니다. (PNG, JPG, JPEG만 가능)'}), 400
    except Exception as e:
        logging.error(f"Image upload failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': f'이미지 업로드 실패: {str(e)}'}), 500

@app.route('/send_email', methods=['POST'])
def send_email():
    participant_ids = request.form.get('participant_ids', '').split(',')
    participant_ids = [int(pid) for pid in participant_ids if pid.strip().isdigit()]
    subject = request.form.get('subject', '')
    body = request.form.get('body', '')
    cc = request.form.get('cc', '')
    include_buttons = request.form.get('include_buttons') == 'yes'
    
    # 세션 데이터 파싱 (행사 프로그램에서 이메일 발송 시)
    session_data_json = request.form.get('session_data', '[]')
    session_data_map = {}  # participant_id -> sessions 배열
    try:
        import json
        if session_data_json and session_data_json != '[]':
            session_data_list = json.loads(session_data_json)
            for item in session_data_list:
                participant_id = int(item.get('participant_id', 0))
                sessions = item.get('sessions', [])
                if participant_id and sessions:
                    session_data_map[participant_id] = sessions
                    logging.info(f"세션 데이터 수신: 참가자 {participant_id}, {len(sessions)}개 세션")
    except Exception as e:
        logging.error(f"세션 데이터 JSON 파싱 오류: {e}")
    
    if not participant_ids or not subject or not body:
        return jsonify({'status': 'error', 'message': 'Missing required fields'}), 400
    
    participants = Participant.query.filter(Participant.id.in_(participant_ids)).all()
    recipient_emails = [p.email for p in participants if p.email]
    if not recipient_emails:
        return jsonify({'status': 'error', 'message': 'No valid recipient emails found.'}), 400
    
    def replace_img_src(match):
        img_tag = match.group(0)
        src_match = re.search(r'src="([^"]+)"', img_tag)
        if src_match:
            img_path = src_match.group(1)
            if img_path.startswith('/uploads/images/'):
                full_path = os.path.join(BASE_DIR, img_path.lstrip('/'))
                if os.path.exists(full_path):
                    with open(full_path, 'rb') as img_file:
                        img_data = img_file.read()
                    return f'<img src="cid:image_{hash(img_path)}" style="max-width: 100%; height: auto;">'
        return img_tag
    
    body = re.sub(r'<img[^>]+>', replace_img_src, body)
    
    # Word 파일들을 루프 밖에서 한 번만 파싱 (모든 참가자에게 동일한 파일 첨부)
    word_files_json = request.form.get('word_files', '[]')
    logging.info(f"🔍 워드 파일 요청 데이터 확인: word_files_json='{word_files_json}' (타입: {type(word_files_json)})")
    logging.info(f"🔍 request.form 전체 키: {list(request.form.keys())}")
    word_files = []
    try:
        import json
        if word_files_json and word_files_json != '[]':
            word_files = json.loads(word_files_json)
            logging.info(f"✅ 워드 파일 정보 수신 성공: {len(word_files)}개 파일")
            for idx, wf in enumerate(word_files):
                logging.info(f"  파일 {idx+1}: name='{wf.get('name')}', server_filename='{wf.get('server_filename')}'")
        else:
            logging.warning(f"⚠️ 워드 파일 정보 없음: word_files_json='{word_files_json}'")
    except Exception as e:
        logging.error(f"❌ 워드 파일 JSON 파싱 오류: {e}, 원본 데이터: {word_files_json}")
        import traceback
        traceback.print_exc()
    
    try:
        for participant in participants:
            if participant.email:
                # 각 참가자별로 실제 동작하는 URL 생성
                accept_url = url_for('accept_response', participant_id=participant.id, response='accept', _external=True)
                decline_url = url_for('accept_response', participant_id=participant.id, response='decline', _external=True)
                
                # 세션 정보 가져오기 (첫 번째 세션 사용)
                session_info = None
                if participant.id in session_data_map and session_data_map[participant.id]:
                    session_info = session_data_map[participant.id][0]  # 첫 번째 세션 사용
                    logging.info(f"세션 정보 로드: 참가자 {participant.id}, 세션 정보: {session_info}")
                else:
                    logging.warning(f"세션 정보 없음: 참가자 {participant.id}")
                
                # Mail merge 처리: *|컬럼명|* 형식의 플레이스홀더를 실제 데이터로 치환
                email_body = apply_mail_merge(body, participant, session_info)
                
                if include_buttons:
                    email_body += '<br><br>'
                    email_body += f'<a href="{accept_url}" style="padding:10px 20px;background:#4CAF50;color:white;text-decoration:none;border-radius:5px;">Accept</a> '
                    email_body += f'<a href="{decline_url}" style="padding:10px 20px;background:#F44336;color:white;text-decoration:none;border-radius:5px;">Decline</a>'
                
                # 제목도 Mail merge 처리
                email_subject = apply_mail_merge(subject, participant, session_info)
                
                # 발신자 설정 - config에서 가져오거나 명시적으로 설정
                sender_email = app.config.get('MAIL_DEFAULT_SENDER') or 'koreaepilepsy@kes.or.kr'
                
                # 수신자 이메일 (참가자 또는 직접 추가한 이메일)
                recipient_email = participant.email if participant else None
                if not recipient_email:
                    # 직접 추가한 이메일인 경우
                    continue  # 이미 recipient_emails에 포함되어 있으므로 다음 루프에서 처리
                
                msg = Message(
                    subject=email_subject,
                    recipients=[recipient_email],
                    html=email_body,
                    sender=sender_email,
                    reply_to='koreaepilepsy@kes.or.kr'
                )
                if cc:
                    msg.cc = [email.strip() for email in cc.split(',') if email.strip()]
                img_pattern = r'<img[^>]+src="([^"]+)"[^>]*>'
                for img_match in re.finditer(img_pattern, email_body):
                    img_path = img_match.group(1)
                    if img_path.startswith('/uploads/images/'):
                        full_path = os.path.join(BASE_DIR, img_path.lstrip('/'))
                        if os.path.exists(full_path):
                            with open(full_path, 'rb') as img_file:
                                msg.attach(
                                    f'image_{hash(img_path)}.png',
                                    'image/png',
                                    img_file.read(),
                                    'inline',
                                    headers=[('Content-ID', f'<image_{hash(img_path)}>')]
                                )
                # Word 파일들을 Mail merge 처리하여 Word 문서로 첨부 (원본 포맷팅 보존)
                # word_files는 루프 밖에서 이미 파싱됨
                
                # Word 파일 첨부 시도 (실패해도 이메일은 발송)
                attachment_errors = []
                if not word_files:
                    logging.warning(f"⚠️ 참가자 {participant.id} ({participant.email}): 첨부할 워드 파일 없음 (word_files 길이: {len(word_files)})")
                else:
                    logging.info(f"📎 참가자 {participant.id} ({participant.email}): {len(word_files)}개 워드 파일 첨부 시도")
                for idx, word_file_info in enumerate(word_files):
                    logging.info(f"📎 워드 파일 {idx+1}/{len(word_files)} 처리 시작")
                    server_filename = word_file_info.get('server_filename')
                    original_filename = word_file_info.get('name', 'document')
                    logging.info(f"워드 파일 처리 시작: server_filename={server_filename}, original_filename={original_filename}, word_file_info={word_file_info}")
                    if server_filename:
                        try:
                            docx_data = generate_mail_merged_docx(server_filename, participant, session_info)
                            if docx_data:
                                # Word 문서를 그대로 첨부 (원본 포맷팅 완벽 보존)
                                # PDF 변환은 포맷팅 손실이 심하므로 Word 문서로 전송
                                participant_name = secure_filename(participant.name_kor or participant.name_eng or participant.first_name or 'participant')
                                file_ext = '.docx' if original_filename.lower().endswith('.docx') else '.doc'
                                docx_filename = f"{participant_name}_{original_filename.replace('.docx', '').replace('.doc', '')}{file_ext}"
                                
                                # MIME 타입 설정
                                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_ext == '.docx' else 'application/msword'
                                
                                # docx_data가 bytes인지 확인
                                if not isinstance(docx_data, bytes):
                                    logging.error(f"docx_data가 bytes가 아닙니다: {type(docx_data)}")
                                    docx_data = bytes(docx_data) if docx_data else None
                                
                                if docx_data:
                                    logging.info(f"📎 Word 문서 첨부 시도: {docx_filename}, 크기: {len(docx_data)} bytes, MIME: {mime_type}")
                                    msg.attach(
                                        docx_filename,
                                        mime_type,
                                        docx_data
                                    )
                                    # 첨부 확인
                                    attachment_count = len(msg.attachments) if hasattr(msg, 'attachments') else 'unknown'
                                    logging.info(f"✅ Word 문서 첨부 완료: {docx_filename} ({len(docx_data)} bytes), 총 첨부 파일 수: {attachment_count}")
                                else:
                                    logging.error(f"❌ docx_data가 None이거나 변환 실패")
                            else:
                                error_msg = f"Word 문서 생성 실패: {server_filename}"
                                logging.warning(error_msg)
                                attachment_errors.append(error_msg)
                        except Exception as e:
                            error_msg = f"워드 파일 처리 오류 ({server_filename}): {str(e)}"
                            logging.error(error_msg)
                            attachment_errors.append(error_msg)
                            import traceback
                            traceback.print_exc()
                    else:
                        error_msg = f"서버 파일명이 없습니다: {word_file_info}"
                        logging.warning(error_msg)
                        attachment_errors.append(error_msg)
                
                # 이메일 발송 전 첨부 파일 확인
                attachment_count = len(msg.attachments) if hasattr(msg, 'attachments') else 0
                logging.info(f"📧 이메일 발송 전 확인 - 수신자: {participant.email}, 첨부 파일 수: {attachment_count}")
                if attachment_count > 0:
                    for att in msg.attachments:
                        logging.info(f"  첨부 파일: {att}")
                
                # 이메일 발송 (PDF 첨부 실패해도 발송)
                try:
                    mail.send(msg)
                    if attachment_errors:
                        logging.warning(f"⚠️ 이메일 발송 성공 (첨부 파일 오류 있음): {participant.email}, 오류: {attachment_errors}")
                    else:
                        logging.info(f"✅ 이메일 발송 성공: {participant.email}, 첨부 파일 수: {attachment_count}")
                except Exception as email_error:
                    logging.error(f"❌ 이메일 발송 실패: {email_error}")
                    import traceback
                    traceback.print_exc()
                    raise
        return jsonify({'status': 'success', 'message': f'Email sent to {len(recipient_emails)} participants'})
    except Exception as e:
        logging.error(f"Email sending failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': f'Failed to send email: {str(e)}'}), 500

@app.route('/upload_word_document', methods=['POST'])
def upload_word_document():
    """워드 문서 업로드 및 서버 저장"""
    try:
        if 'word_document' not in request.files:
            return jsonify({'status': 'error', 'message': '파일이 없습니다.'}), 400
        
        file = request.files['word_document']
        if file.filename == '':
            return jsonify({'status': 'error', 'message': '파일이 선택되지 않았습니다.'}), 400
        
        # 파일 확장자 확인
        if not (file.filename.endswith('.doc') or file.filename.endswith('.docx')):
            return jsonify({'status': 'error', 'message': '워드 문서 파일만 업로드 가능합니다.'}), 400
        
        # 고유한 파일명 생성 (타임스탬프 + UUID)
        original_filename = secure_filename(file.filename)
        file_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        server_filename = f"{timestamp}_{file_id}_{original_filename}"
        filepath = os.path.join(app.config['WORD_UPLOAD_FOLDER'], server_filename)
        
        # 파일 저장
        file.save(filepath)
        
        return jsonify({
            'status': 'success',
            'server_filename': server_filename,
            'file_id': file_id,
            'original_filename': original_filename,
            'message': '워드 문서가 업로드되었습니다.'
        })
    except Exception as e:
        logging.error(f"Word document upload failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': f'워드 문서 업로드 실패: {str(e)}'}), 500

def convert_word_to_html(docx_path):
    """워드 문서를 HTML로 변환"""
    try:
        # python-docx 라이브러리 사용 시도
        try:
            from docx import Document
            doc = Document(docx_path)
            html_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # 단락 스타일 확인
                    style = para.style.name if para.style else 'Normal'
                    text = para.text
                    
                    # 텍스트 포맷팅 처리
                    runs_html = []
                    for run in para.runs:
                        run_text = run.text
                        if not run_text:
                            continue
                        
                        # 포맷팅 적용
                        if run.bold:
                            run_text = f'<strong>{run_text}</strong>'
                        if run.italic:
                            run_text = f'<em>{run_text}</em>'
                        if run.underline:
                            run_text = f'<u>{run_text}</u>'
                        
                        # 색상 처리
                        if run.font.color and run.font.color.rgb:
                            color = f"#{run.font.color.rgb:06x}"
                            run_text = f'<span style="color: {color}">{run_text}</span>'
                        
                        runs_html.append(run_text)
                    
                    para_html = ''.join(runs_html) if runs_html else text
                    
                    # 스타일에 따라 태그 추가
                    if style.startswith('Heading'):
                        level = style.replace('Heading ', '').replace('Heading', '1')
                        try:
                            level = int(level) if level.isdigit() else 1
                        except:
                            level = 1
                        html_parts.append(f'<h{level}>{para_html}</h{level}>')
                    else:
                        html_parts.append(f'<p>{para_html}</p>')
                
                # 빈 줄 추가
                if not para.text.strip() and html_parts:
                    html_parts.append('<br>')
            
            return '\n'.join(html_parts) if html_parts else '<p>문서 내용이 없습니다.</p>'
        except ImportError:
            # python-docx가 설치되지 않은 경우 기본 처리
            logging.warning("python-docx가 설치되지 않았습니다. 기본 변환을 사용합니다.")
            return '<p>워드 문서 변환을 위해 python-docx 라이브러리가 필요합니다. pip install python-docx로 설치해주세요.</p>'
    except Exception as e:
        logging.error(f"Word to HTML conversion failed: {e}")
        return f'<p>워드 문서 변환 중 오류가 발생했습니다: {str(e)}</p>'

def get_participant_data_dict(participant, session_info=None):
    """참가자 데이터를 딕셔너리로 반환 (한글/영어 컬럼명 모두 지원, 세션 정보 포함)"""
    if not participant:
        return {}
    
    # 기본 데이터 딕셔너리
    data = {
        'name': participant.name_kor or participant.name_eng or (participant.first_name + ' ' + participant.family_name if participant.first_name and participant.family_name else ''),
        'name_kor': participant.name_kor or '',
        'name_eng': participant.name_eng or '',
        'first_name': participant.first_name or '',
        'family_name': participant.family_name or '',
        'email': participant.email or '',
        'affiliation_kor': participant.affiliation_kor or '',
        'affiliation_eng': participant.affiliation_eng or '',
        'country': participant.country or '',
        'country_code': participant.country_code or '',
        'phone': participant.phone or '',
        'position': participant.position or '',
        'role': participant.role or '',
        'code': participant.code or '',
        'department_eng': participant.department_eng or ''
    }
    
    # 세션 정보 추가 (행사 프로그램에서 이메일 발송 시)
    if session_info:
        data.update({
            'date': session_info.get('date', ''),
            'session_type': session_info.get('session_type', ''),
            'language': session_info.get('language', ''),
            'session_abbreviation': session_info.get('session_abbreviation', ''),
            'session_title': session_info.get('session_title', ''),
            'venue': session_info.get('venue', ''),
            'session_time': session_info.get('session_time', ''),
            'session_start_time': session_info.get('session_start_time', ''),
            'session_end_time': session_info.get('session_end_time', ''),
            'lecture_title': session_info.get('lecture_title', ''),
            'lecture_time': session_info.get('lecture_time', ''),
            'lecture_start_time': session_info.get('lecture_start_time', ''),
            'lecture_end_time': session_info.get('lecture_end_time', '')
        })
    
    # 한글 컬럼명 매핑 추가
    korean_mapping = {
        '이메일': 'email',
        '이름': 'name',
        '이름(한글)': 'name_kor',
        '이름(영문)': 'name_eng',
        '이름한글': 'name_kor',
        '이름영문': 'name_eng',
        '성': 'family_name',
        '이름(영어)': 'first_name',
        '소속': 'affiliation_kor',
        '소속(한글)': 'affiliation_kor',
        '소속(영문)': 'affiliation_eng',
        '소속한글': 'affiliation_kor',
        '소속영문': 'affiliation_eng',
        '국가': 'country',
        '전화': 'phone',
        '전화번호': 'phone',
        '직위': 'position',
        '역할': 'role',
        '코드': 'code',
        '과(영문)': 'department_eng',
        '과영문': 'department_eng',
        # 세션 관련 한글 매핑
        '날짜': 'date',
        '세션종류': 'session_type',
        '세션 종류': 'session_type',
        '언어': 'language',
        '세션약어': 'session_abbreviation',
        '세션 약어': 'session_abbreviation',
        '세션명': 'session_title',
        '세션 주제': 'session_title',
        '세션주제': 'session_title',
        '장소': 'venue',
        '발표장소': 'venue',
        '세션시간': 'session_time',
        '세션 시간': 'session_time',
        '발표시간': 'lecture_time',
        '발표 시간': 'lecture_time',
        '발표주제': 'lecture_title',
        '발표 주제': 'lecture_title'
    }
    
    # 한글 컬럼명도 추가
    for korean_name, english_key in korean_mapping.items():
        if english_key in data:
            data[korean_name] = data[english_key]
    
    return data

def apply_mail_merge(text, participant, session_info=None):
    """Mail merge 플레이스홀더를 실제 데이터로 치환"""
    if not text or not participant:
        return text
    
    participant_data = get_participant_data_dict(participant, session_info)
    
    # *|컬럼명|* 형식의 플레이스홀더 찾아서 치환
    def replace_placeholder(match):
        column_name = match.group(1).strip()
        # 원본 컬럼명 (한글 지원)
        original_name = column_name
        # 소문자 변환 (영어 컬럼명 매칭용)
        column_name_lower = column_name.lower().replace('-', '_')
        
        # 먼저 원본 이름으로 찾기 (한글 컬럼명 지원)
        if original_name in participant_data:
            value = participant_data[original_name]
            return str(value) if value else ''
        
        # 소문자로 찾기 (영어 컬럼명 지원)
        if column_name_lower in participant_data:
            value = participant_data[column_name_lower]
            return str(value) if value else ''
        
        # 찾을 수 없으면 원본 반환
        return match.group(0)
    
    # *|컬럼명|* 패턴 찾기
    pattern = r'\*\|\s*([^|]+)\s*\|\*'
    result = re.sub(pattern, replace_placeholder, text)
    
    return result

def generate_mail_merged_docx(server_filename, participant, session_info=None):
    """Word 문서를 Mail merge 처리하여 Word 문서로 반환 (원본 포맷팅 보존)"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        
        word_filepath = os.path.join(app.config['WORD_UPLOAD_FOLDER'], server_filename)
        logging.info(f"🔍 워드 파일 경로 확인: {word_filepath}")
        logging.info(f"🔍 WORD_UPLOAD_FOLDER: {app.config['WORD_UPLOAD_FOLDER']}")
        logging.info(f"🔍 server_filename: {server_filename}")
        logging.info(f"🔍 파일 존재 여부: {os.path.exists(word_filepath)}")
        if not os.path.exists(word_filepath):
            logging.error(f"❌ Word file not found: {word_filepath}")
            # 절대 경로로도 확인
            abs_path = os.path.abspath(word_filepath)
            logging.error(f"❌ 절대 경로: {abs_path}, 존재 여부: {os.path.exists(abs_path)}")
            return None
        
        # Word 문서 열기
        doc = Document(word_filepath)
        participant_data = get_participant_data_dict(participant, session_info)
        
        # 디버깅: 사용 가능한 컬럼명 로깅
        logging.info(f"Mail merge - 사용 가능한 컬럼명: {list(participant_data.keys())}")
        logging.info(f"Mail merge - participant_data 샘플: {dict(list(participant_data.items())[:10])}")
        logging.info(f"Mail merge - session_info: {session_info}")
        if session_info:
            logging.info(f"Mail merge - 세션 데이터 확인: date={session_info.get('date')}, session_title={session_info.get('session_title')}, venue={session_info.get('venue')}, lecture_title={session_info.get('lecture_title')}, lecture_time={session_info.get('lecture_time')}")
        
        # Mail merge 처리: 모든 단락과 테이블에서 플레이스홀더 치환
        def replace_placeholder_in_text(text):
            if not text:
                return text
            
            def replace_placeholder(match):
                column_name = match.group(1).strip()
                # 원본 컬럼명 (한글 지원)
                original_name = column_name
                # 소문자 변환 (영어 컬럼명 매칭용)
                column_name_lower = column_name.lower().replace('-', '_')
                
                # 먼저 원본 이름으로 찾기 (한글 컬럼명 지원)
                if original_name in participant_data:
                    value = participant_data[original_name]
                    logging.info(f"✅ 플레이스홀더 치환: *|{column_name}|* -> '{value}'")
                    return str(value) if value else ''
                
                # 소문자로 찾기 (영어 컬럼명 지원)
                if column_name_lower in participant_data:
                    value = participant_data[column_name_lower]
                    logging.info(f"✅ 플레이스홀더 치환 (소문자): *|{column_name}|* -> '{value}'")
                    return str(value) if value else ''
                
                # 찾을 수 없으면 경고 로그와 함께 원본 반환
                available_keys = [k for k in participant_data.keys() if participant_data.get(k)]
                logging.warning(f"⚠️ 플레이스홀더를 찾을 수 없음: *|{column_name}|* (사용 가능한 키: {available_keys[:30]})")
                return match.group(0)
            
            pattern = r'\*\|\s*([^|]+)\s*\|\*'
            return re.sub(pattern, replace_placeholder, text)
        
        # 단락에서 플레이스홀더 치환 (run별로 처리하여 포맷팅과 이미지 보존)
        def replace_placeholders_in_paragraph(para):
            """단락의 플레이스홀더를 치환하되 포맷팅과 이미지를 보존"""
            if not para.text:
                return
            
            # 단락 전체 텍스트를 합치기
            full_text = ''.join([run.text for run in para.runs if run.text])
            
            if not full_text:
                return
            
            # 플레이스홀더가 있는지 확인
            if '*|' not in full_text or '|*' not in full_text:
                return  # 플레이스홀더가 없으면 건너뛰기
            
            # 단락 전체 텍스트에서 플레이스홀더 치환
            replaced_text = replace_placeholder_in_text(full_text)
            
            if full_text == replaced_text:
                return  # 변경사항 없음
            
            logging.info(f"플레이스홀더 치환: '{full_text[:150]}...' -> '{replaced_text[:150]}...'")
            
            # 각 run에서 직접 플레이스홀더 치환 (포맷팅 보존)
            # 이 방법이 색상을 포함한 모든 포맷팅을 보존하는 가장 확실한 방법
            for run in para.runs:
                if run.text and '*|' in run.text:
                    # run의 텍스트에서 플레이스홀더 치환
                    original_run_text = run.text
                    replaced_run_text = replace_placeholder_in_text(original_run_text)
                    
                    if original_run_text != replaced_run_text:
                        # 텍스트만 교체 (포맷팅은 그대로 유지됨)
                        run.text = replaced_run_text
                        logging.debug(f"Run 텍스트 치환: '{original_run_text}' -> '{replaced_run_text}'")
            
            # 플레이스홀더가 여러 run에 걸쳐 있는 경우를 처리
            # 단락 전체 텍스트를 다시 확인하여 치환이 완료되었는지 확인
            remaining_text = ''.join([run.text for run in para.runs if run.text])
            if '*|' in remaining_text and remaining_text != replaced_text:
                # 여러 run에 걸쳐 있는 경우, 단락을 재구성
                # 각 run의 포맷팅 정보를 저장
                runs_info = []
                for run in para.runs:
                    if run.text:
                        runs_info.append({
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font_name': run.font.name if run.font else None,
                            'font_size': run.font.size.pt if run.font and run.font.size else None,
                            'font_color': run.font.color.rgb if run.font and run.font.color else None
                        })
                
                # 단락을 clear하고 치환된 텍스트로 재구성
                para.clear()
                
                # 첫 번째 run의 포맷팅을 사용하여 새 run 생성
                if runs_info:
                    first_run = runs_info[0]
                    new_run = para.add_run(replaced_text)
                    
                    # 포맷팅 복원
                    if first_run.get('bold') is not None:
                        new_run.bold = first_run['bold']
                    if first_run.get('italic') is not None:
                        new_run.italic = first_run['italic']
                    if first_run.get('underline') is not None:
                        new_run.underline = first_run['underline']
                    if first_run.get('font_name'):
                        new_run.font.name = first_run['font_name']
                    if first_run.get('font_size'):
                        new_run.font.size = Pt(first_run['font_size'])
                    
                    # 색상 복원 - RGBColor 객체를 직접 할당
                    if first_run.get('font_color') is not None:
                        try:
                            rgb_obj = first_run['font_color']
                            new_run.font.color.rgb = rgb_obj
                            logging.info(f"✅ 색상 복원 (여러 run에 걸친 경우): {rgb_obj}")
                        except Exception as e:
                            logging.warning(f"⚠️ 색상 복원 실패: {e}")
                else:
                    para.add_run(replaced_text)
        
        # 모든 단락 처리
        paragraph_count = 0
        replaced_count = 0
        for para in doc.paragraphs:
            paragraph_count += 1
            original_text = para.text
            if original_text and '*|' in original_text:
                replace_placeholders_in_paragraph(para)
                if para.text != original_text:
                    replaced_count += 1
                    logging.info(f"단락 {paragraph_count} 치환 완료: '{original_text[:80]}...' -> '{para.text[:80]}...'")
        
        logging.info(f"총 {paragraph_count}개 단락 중 {replaced_count}개 단락에서 플레이스홀더 치환됨")
        
        # 테이블에서 플레이스홀더 치환
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        if original_text and '*|' in original_text:
                            replace_placeholders_in_paragraph(para)
                            if para.text != original_text:
                                logging.info(f"테이블 {table_count} 셀 치환 완료: '{original_text[:80]}...' -> '{para.text[:80]}...'")
        
        logging.info(f"총 {table_count}개 테이블 처리 완료")
        
        # Word 문서를 메모리에 저장
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_data = buffer.read()
        buffer.close()
        
        return docx_data
        
    except Exception as e:
        logging.error(f"Word 문서 생성 실패: {e}")
        import traceback
        traceback.print_exc()
        return None

def convert_docx_to_pdf_from_bytes(docx_bytes):
    """Word 문서 바이트를 PDF로 변환 (원본 포맷팅 보존)"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
        
        # 바이트에서 Word 문서 로드
        doc = Document(BytesIO(docx_bytes))
        
        # 기존 함수 사용
        return convert_docx_to_pdf_from_doc(doc)
        
    except Exception as e:
        logging.error(f"바이트에서 PDF 변환 실패: {e}")
        import traceback
        traceback.print_exc()
        return None

def convert_docx_to_pdf_from_doc(doc):
    """워드 문서 객체를 PDF로 변환 (원본 포맷팅 최대한 보존 - reportlab 개선 버전)"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from io import BytesIO
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import platform
        
        # PDF를 메모리에 생성
        buffer = BytesIO()
        
        # A4 크기로 PDF 생성
        pdf_doc = SimpleDocTemplate(buffer, pagesize=A4,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=72)
        
        # 한글 폰트 등록
        korean_font_name = 'KoreanFont'
        korean_font_registered = False
        
        # macOS에서 한글 폰트 찾기
        if platform.system() == 'Darwin':  # macOS
            font_paths = [
                '/System/Library/Fonts/Supplemental/AppleGothic.ttf',
                '/Library/Fonts/AppleGothic.ttf',
                '/System/Library/Fonts/AppleGothic.ttf',
                '/System/Library/Fonts/Supplemental/AppleSDGothicNeo-Regular.ttf',
            ]
        else:
            # Linux/Windows에서 한글 폰트 찾기
            font_paths = [
                '/usr/share/fonts/truetype/nanum/NanumGothic.ttf',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                'C:/Windows/Fonts/malgun.ttf',  # Windows 맑은 고딕
                'C:/Windows/Fonts/gulim.ttc',   # Windows 굴림
            ]
        
        for font_path in font_paths:
            try:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont(korean_font_name, font_path))
                    korean_font_registered = True
                    logging.info(f"한글 폰트 등록 성공: {font_path}")
                    break
            except Exception as e:
                logging.warning(f"폰트 등록 실패 ({font_path}): {e}")
                continue
        
        if not korean_font_registered:
            logging.warning("한글 폰트를 찾을 수 없습니다. 기본 폰트를 사용합니다. 한글이 제대로 표시되지 않을 수 있습니다.")
            korean_font_name = 'Helvetica'  # 폴백
        
        # 스타일 정의
        styles = getSampleStyleSheet()
        story = []
        
        # 기본 스타일 설정 (한글 폰트 사용)
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=korean_font_name,
            fontSize=12,
            leading=14,
            alignment=TA_LEFT,
        )
        
        # 단락 처리
        for para in doc.paragraphs:
            if para.text.strip():
                # 텍스트 포맷팅 처리
                para_text = para.text
                
                # HTML 이스케이프 (한글 텍스트 보존)
                para_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # 스타일에 따라 처리
                style_name = para.style.name if para.style else 'Normal'
                
                # Bold, Italic, Underline 처리 (run별로 정확하게 처리)
                formatted_parts = []
                for run in para.runs:
                    if not run.text:
                        continue
                    
                    run_text = run.text
                    # HTML 이스케이프 (한글 텍스트 보존)
                    run_text = run_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    # 포맷팅 태그 적용
                    if run.bold:
                        run_text = f'<b>{run_text}</b>'
                    if run.italic:
                        run_text = f'<i>{run_text}</i>'
                    if run.underline:
                        run_text = f'<u>{run_text}</u>'
                    
                    formatted_parts.append(run_text)
                
                formatted_text = ''.join(formatted_parts) if formatted_parts else para_text
                
                # UTF-8 인코딩 확인 및 로깅 (디버깅용)
                try:
                    formatted_text.encode('utf-8')
                except UnicodeEncodeError as e:
                    logging.error(f"인코딩 오류: {e}, 텍스트: {formatted_text[:50]}")
                
                if style_name.startswith('Heading'):
                    level = style_name.replace('Heading ', '').replace('Heading', '1')
                    try:
                        level = int(level) if level.isdigit() else 1
                        level = max(1, min(6, level))
                        font_size = 18 - (level - 1) * 2
                        heading_style = ParagraphStyle(
                            f'Heading{level}',
                            parent=styles['Heading1'],
                            fontName=korean_font_name,
                            fontSize=font_size,
                            leading=font_size + 4,
                            spaceAfter=12,
                        )
                        story.append(Paragraph(formatted_text, heading_style))
                    except:
                        fallback_heading = ParagraphStyle(
                            'HeadingFallback',
                            parent=styles['Heading1'],
                            fontName=korean_font_name,
                        )
                        story.append(Paragraph(formatted_text, fallback_heading))
                else:
                    story.append(Paragraph(formatted_text, normal_style))
                
                story.append(Spacer(1, 6))
            else:
                # 빈 줄
                story.append(Spacer(1, 6))
        
        # 테이블 처리
        for table in doc.tables:
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_text.append(para.text)
                    cell_content = ' '.join(cell_text) if cell_text else ''
                    # HTML 이스케이프
                    cell_content = cell_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    row_data.append(Paragraph(cell_content, normal_style))
                table_data.append(row_data)
            
            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 12))
        
        # PDF 빌드
        pdf_doc.build(story)
        
        # PDF 바이트 반환
        buffer.seek(0)
        pdf_bytes = buffer.read()
        buffer.close()
        
        return pdf_bytes
        
    except ImportError:
        logging.error("reportlab이 설치되지 않았습니다. pip install reportlab로 설치해주세요.")
        return None
    except Exception as e:
        logging.error(f"reportlab PDF 변환 실패: {e}")
        import traceback
        traceback.print_exc()
        return None

def convert_docx_to_pdf(docx_path):
    """Word 문서 파일을 PDF로 변환 (순수 Python 라이브러리만 사용 - 권한 문제 없음)"""
    try:
        from docx import Document
        
        # 워드 문서 열기
        doc = Document(docx_path)
        
        # 문서 객체를 PDF로 변환
        return convert_docx_to_pdf_from_doc(doc)
        
    except Exception as e:
        logging.error(f"PDF 변환 오류: {e}")
        import traceback
        traceback.print_exc()
        return None


@app.route('/uploads/images/<filename>')
def serve_image(filename):
    """이미지 파일 서빙"""
    return send_from_directory(app.config['IMAGE_UPLOAD_FOLDER'], filename)

@app.route('/decline_popup', methods=['GET'])
def decline_popup():
    """거절 사유 팝업"""
    participant_id = request.args.get('participant_id')
    return render_template('decline_reason.html', participant_id=participant_id)

@app.route('/accept_response', methods=['GET', 'POST'])
def accept_response():
    """참가자 응답 처리"""
    if request.method == 'GET':
        participant_id = request.args.get('participant_id')
        response = request.args.get('response')  # 'accept' or 'decline'
        participant = Participant.query.get(participant_id) if participant_id else None
        
        # Accept인 경우 즉시 처리
        if response == 'accept' and participant:
            participant.accept_or_decline = 'Accept'
            participant.decline_reason = None
            db.session.commit()
        
        return render_template('accept_response.html', participant=participant, response=response)
    
    elif request.method == 'POST':
        participant_id = request.form.get('participant_id')
        response = request.form.get('response')
        decline_reason = request.form.get('decline_reason', '')
        
        participant = Participant.query.get_or_404(participant_id)
        
        if response == 'accept':
            participant.accept_or_decline = 'Accept'
            participant.decline_reason = None
        elif response == 'decline':
            participant.accept_or_decline = 'Decline'
            participant.decline_reason = decline_reason
        
        db.session.commit()
        
        return jsonify({'success': True})

@app.route('/view_decline_reason/<int:participant_id>')
def view_decline_reason(participant_id):
    """거절 사유 조회"""
    participant = Participant.query.get_or_404(participant_id)
    return jsonify({'decline_reason': participant.decline_reason or ''})

@app.route('/decline_reason')
def decline_reason():
    """거절 사유 표시 페이지"""
    participant_name = request.args.get('participant_name', '')
    reason = request.args.get('reason', '')
    
    # URL 디코딩 처리
    import urllib.parse
    participant_name = urllib.parse.unquote(participant_name)
    reason = urllib.parse.unquote(reason)
    
    return render_template('decline_reason_display.html', participant_name=participant_name, reason=reason)

@app.route('/participant_by_code', methods=['GET'])
def get_participant_by_code():
    code = request.args.get('code')
    event_id = request.args.get('event_id')
    if not code or not event_id:
        return jsonify({'status': 'error', 'message': 'Code and event_id are required'}), 400
    try:
        code = int(code)
        event_id = int(event_id)
    except Exception:
        return jsonify({'status': 'error', 'message': 'Invalid code or event_id'}), 400
    participant = Participant.query.filter_by(code=code, event_id=event_id).first()
    if not participant:
        return jsonify({'status': 'error', 'message': '참가자를 찾을 수 없습니다.'}), 404
    return jsonify({
        'status': 'success',
        'participant': {
        'id': participant.id,
            'name_kor': participant.name_kor or f"{participant.first_name} {participant.family_name}",
        'event_id': participant.event_id,
            'accept_or_decline': participant.accept_or_decline,
            'check_in_time': participant.check_in_time.strftime('%Y-%m-%d %H:%M:%S') if participant.check_in_time else None,
            'check_out_time': participant.check_out_time.strftime('%Y-%m-%d %H:%M:%S') if participant.check_out_time else None
        }
    })

@app.route('/participant/check_attendance_by_code', methods=['POST'])
def check_attendance_by_code():
    """통합 출석 처리: 첫 번째는 체크인, 그 다음부터는 체크아웃"""
    code = request.form.get('code') or request.args.get('code')
    event_id = request.form.get('event_id') or request.args.get('event_id')
    if not code or not event_id:
        return jsonify({'status': 'error', 'message': 'Code and event_id are required'}), 400
    try:
        code = int(code)
        event_id = int(event_id)
    except Exception:
        return jsonify({'status': 'error', 'message': 'Invalid code or event_id'}), 400
    
    participant = Participant.query.filter_by(code=code, event_id=event_id).first()
    if not participant:
        return jsonify({'status': 'error', 'message': 'Participant not found'}), 404
    
    from datetime import datetime, timezone, timedelta
    
    # 서버 환경 자동 감지 및 한국 시간 계산
    import os
    import time
    
    # 시스템 시간대 확인
    try:
        # 시스템 시간대 정보 가져오기
        timezone_info = time.tzname[time.daylight]
        print(f"DEBUG: System timezone: {timezone_info}")
        
        # 환경변수로 시간대 확인
        env_tz = os.environ.get('TZ', '')
        print(f"DEBUG: Environment TZ: {env_tz}")
        
        # 현재 로컬 시간과 UTC 시간 비교
        local_time = datetime.now()
        utc_time = datetime.now(timezone.utc)
        
        print(f"DEBUG: Local time: {local_time}")
        print(f"DEBUG: UTC time: {utc_time}")
        print(f"DEBUG: Time difference: {local_time.hour - utc_time.hour} hours")
        
        # 한국 시간대인지 확인 (KST, Asia/Seoul, +09:00 등)
        is_korea_timezone = (
            'KST' in str(timezone_info).upper() or 
            'KOREA' in str(timezone_info).upper() or 
            'SEOUL' in str(timezone_info).upper() or
            'ASIA/SEOUL' in str(env_tz).upper() or
            '+09' in str(env_tz)
        )
        
        # UTC 환경인지 명시적으로 확인
        is_utc_timezone = (
            'UTC' in str(timezone_info).upper() or
            'ETC/UTC' in str(env_tz).upper() or
            env_tz == '' or  # TZ 환경변수가 설정되지 않음
            abs(local_time.hour - utc_time.hour) <= 1  # 로컬과 UTC 시간이 같거나 1시간 이내 차이
        )
        
        if is_korea_timezone:
            # 이미 한국 시간대
            now = local_time
            print(f"DEBUG: Server is in Korea timezone, using local time")
        elif is_utc_timezone:
            # UTC 시간대이므로 +9시간 추가
            now = utc_time.replace(tzinfo=timezone.utc) + timedelta(hours=9)
            print(f"DEBUG: Server is in UTC timezone, converting to Korea time (+9 hours)")
        else:
            # 기타 시간대는 UTC + 9시간으로 처리
            now = utc_time.replace(tzinfo=timezone.utc) + timedelta(hours=9)
            print(f"DEBUG: Server timezone unknown, using UTC + 9 hours as fallback")
            
    except Exception as e:
        # 에러 발생 시 기본적으로 UTC + 9시간 사용
        print(f"DEBUG: Error detecting timezone: {e}")
        now = datetime.now(timezone.utc) + timedelta(hours=9)
        print(f"DEBUG: Using fallback: UTC + 9 hours")
    
    # 디버깅 로그
    print(f"DEBUG: Code {code}, Event {event_id}")
    print(f"DEBUG: Current check_in_time: {participant.check_in_time}")
    print(f"DEBUG: Current check_out_time: {participant.check_out_time}")
    print(f"DEBUG: Final Korea time: {now}")
    print(f"DEBUG: Final Korea time strftime: {now.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # 참가자 이름
    participant_name = participant.name_kor or f"{participant.first_name} {participant.family_name}"
    
    # 체크인이 없으면 체크인으로 처리
    if not participant.check_in_time:
        print(f"DEBUG: Performing CHECK-IN for participant {participant_name}")
        participant.check_in_time = now
        # 체크인 시에는 체크아웃 시간을 명시적으로 None으로 설정
        participant.check_out_time = None
        db.session.commit()
        print(f"DEBUG: After check-in - check_in_time: {participant.check_in_time}, check_out_time: {participant.check_out_time}")
        # 서버 시간을 그대로 사용 (이미 한국 시간)
        korea_time_str = participant.check_in_time.strftime('%Y-%m-%d %H:%M:%S')
        
        return jsonify({
            'status': 'success', 
            'message': 'Check-in successful', 
            'action': 'check_in',
            'participant_name': participant_name,
            'check_in_time': korea_time_str,  # 서버 로컬 시간 (한국 시간)
            'check_out_time': None  # 체크인 시에는 체크아웃 시간을 None으로 반환
        })
    
    # 체크인이 있으면 체크아웃으로 처리
    else:
        print(f"DEBUG: Performing CHECK-OUT for participant {participant_name}")
        participant.check_out_time = now
        db.session.commit()
        print(f"DEBUG: After check-out - check_in_time: {participant.check_in_time}, check_out_time: {participant.check_out_time}")
        # 서버 시간을 그대로 사용 (이미 한국 시간으로 변환됨)
        korea_checkin_str = participant.check_in_time.strftime('%Y-%m-%d %H:%M:%S')
        korea_checkout_str = participant.check_out_time.strftime('%Y-%m-%d %H:%M:%S')
        
        return jsonify({
            'status': 'success', 
            'message': 'Check-out successful', 
            'action': 'check_out',
            'participant_name': participant_name,
            'check_in_time': korea_checkin_str,  # 서버 로컬 시간 (한국 시간)
            'check_out_time': korea_checkout_str  # 서버 로컬 시간 (한국 시간)
        })

@app.route('/track_checkin/<int:event_id>')
def track_checkin(event_id):
    """출석 관리 페이지"""
    event = Event.query.get_or_404(event_id)
    return render_template('track_checkin.html', event=event)

@app.route('/download_participants_excel/<int:event_id>', methods=['POST'])
def download_participants_excel(event_id):
    """참가자 엑셀 다운로드"""
    try:
        print(f"=== EXCEL DOWNLOAD REQUEST ===")
        print(f"Event ID: {event_id}")
        print(f"Form data: {dict(request.form)}")
        
        event = Event.query.get_or_404(event_id)
        print(f"Event found: {event.name}")
        
        selected_columns = request.form.getlist('selected_columns')
        participant_ids = request.form.get('selected_participants', '')
        
        print(f"Selected columns: {selected_columns}")
        print(f"Participant IDs: {participant_ids}")
        
        if participant_ids:
            ids = [int(pid) for pid in participant_ids.split(',') if pid.strip().isdigit()]
            participants = Participant.query.filter(Participant.id.in_(ids)).all()
            print(f"Filtered participants by IDs: {len(participants)}")
        else:
            participants = Participant.query.filter_by(event_id=event_id).all()
            print(f"All participants for event: {len(participants)}")
        
        if not participants:
            print("No participants found")
            return "No participants found", 400
        
        # 코드 번호로 오름차순 정렬
        participants = sorted(participants, key=lambda p: int(p.code) if p.code and (isinstance(p.code, int) or (isinstance(p.code, str) and p.code.isdigit())) else 0)
        
        # 데이터 준비 (새로운 컬럼 순서)
        data = []
        for p in participants:
            # 성명(ENG) 자동 조합 로직
            name_eng_display = p.name_eng
            if not name_eng_display:
                if p.first_name and p.family_name:
                    name_eng_display = f"{p.first_name} {p.family_name}"
                elif p.first_name:
                    name_eng_display = p.first_name
                elif p.family_name:
                    name_eng_display = p.family_name
                else:
                    name_eng_display = ''
            
            data.append({
                '이벤트 ID': event.event_id,
                '코드': p.code,
                '역할': p.role,
                '국가': p.country,
                '국가약어': p.country_code,
                '성명(KOR)': p.name_kor,
                '성명(ENG)': name_eng_display,
                '이름(First Name)': p.first_name,
                '성(Last Name)': p.family_name,
                '이메일': p.email,
                '전화': p.phone,
                '소속(ENG)': p.affiliation_eng,
                '과(ENG)': p.department_eng,
                '소속(KOR)': p.affiliation_kor,
                '과(KOR)': p.department_kor,
                '직위': p.position,
                '면허번호': p.license_number,
                '생년월일': p.birth_date.strftime('%Y-%m-%d') if p.birth_date else '',
                '회원구분': p.workplace_type,
                '등록구분': p.registration,
                '승인/거절': p.accept_or_decline,
                'CV': p.cv,
                '사진': p.photo,
                'PPT': p.ppt,
                'Script': p.script,
                '동의여부': p.agree,
                '비고(사용자)': p.remark_user,
                '비고(관리자)': p.remark_admin,
                '체크인': p.check_in_time.strftime('%Y-%m-%d %H:%M:%S') if p.check_in_time else '',
                '체크아웃': p.check_out_time.strftime('%Y-%m-%d %H:%M:%S') if p.check_out_time else '',
                '거절 사유': p.decline_reason
            })
        
        print(f"Data prepared: {len(data)} rows")
        
        # 엑셀 파일 생성
        df = pd.DataFrame(data)
        if selected_columns:
            # 선택된 컬럼만 필터링
            available_columns = [col for col in selected_columns if col in df.columns]
            if available_columns:
                df = df[available_columns]
            else:
                print("Warning: No valid columns selected, using all columns")
        
        print(f"DataFrame shape: {df.shape}")
        
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name='Participants', index=False)
        
        output.seek(0)
        file_size = len(output.getvalue())
        print(f"Excel file created, size: {file_size} bytes")
        
        if file_size == 0:
            print("Error: Generated file is empty")
            return "Generated file is empty", 500
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'participants_{event.event_id}.xlsx'
        )
        
    except Exception as e:
        print(f"Error in download_participants_excel: {e}")
        import traceback
        traceback.print_exc()
        return f"Error generating Excel file: {str(e)}", 500

@app.route('/get_participant_status', methods=['GET'])
def get_participant_status():
    event_id = request.args.get('event_id', type=int)
    if event_id:
        participants = Participant.query.filter_by(event_id=event_id).all()
    else:
        participants = Participant.query.all()
    def format_time(dt):
        if not dt:
            return ''
        # datetime 객체면 strftime, 아니면 문자열로 처리
        if hasattr(dt, 'strftime'):
            return dt.strftime('%Y-%m-%d %H:%M')
        s = str(dt)
        # YYYY-MM-DD HH:MM:SS(.마이크로초) → YYYY-MM-DD HH:MM
        if len(s) >= 16:
            return s[:16]
        return s
    return jsonify([{
        'event_id': p.event_id,
        'code': p.code,
        'accept_or_decline': p.accept_or_decline,
        'check_in_time': format_time(p.check_in_time),
        'check_out_time': format_time(p.check_out_time),
        'remark_user': p.remark_user,
        'remark_admin': p.remark_admin,
        'decline_reason': p.decline_reason
    } for p in participants])

@app.route('/select_columns')
def select_columns():
    event_id = request.args.get('event_id')
    participants = request.args.get('participants', '')
    
    print(f"=== SELECT COLUMNS REQUEST ===")
    print(f"Event ID: {event_id}")
    print(f"Participants: {participants}")
    print(f"All args: {dict(request.args)}")
    
    columns = [
        'Event ID', 'Code', 'Registration', 'Division', 'Role', 'Country',
        'Name (KOR)', 'Affiliation (KOR)', 'Department (KOR)', 'First Name', 'Family Name',
        'Affiliation (ENG)', 'Department (ENG)', 'Accept/Decline', 'Email', 'Phone',
        'Position', 'License #', 'Agree', 'Remark (User)', 'Remark (Admin)',
        'Check-In Time', 'Check-Out Time', 'Decline Reason'
    ]
    return render_template('select_columns.html', event_id=event_id, participants=participants, columns=columns)

# 참가자 등록 페이지 (관리자용)
@app.route('/register/<int:event_id>')
def participant_registration(event_id):
    """참가자 등록 페이지"""
    event = Event.query.get_or_404(event_id)
    return render_template('participant_registration.html', event=event, timedelta=timedelta)

# 공개 참가자 등록 페이지
@app.route('/public/register/<int:event_id>')
def public_participant_registration(event_id):
    """공개 참가자 등록 페이지"""
    event = Event.query.get_or_404(event_id)
    return render_template('participant_registration.html', event=event, timedelta=timedelta, is_public=True)

# 참가자 등록 처리
@app.route('/register_participant/<int:event_id>', methods=['POST'])
def register_participant(event_id):
    """참가자 등록 처리"""
    try:
        event = Event.query.get_or_404(event_id)
        
        # 폼 데이터 수집
        registration_type = request.form.get('registration_type')
        name_kor = request.form.get('name_kor')
        first_name = request.form.get('first_name')
        family_name = request.form.get('family_name')
        license_number = request.form.get('license_number')
        specialty = request.form.get('specialty')
        specialty_other = request.form.get('specialty_other')
        phone = request.form.get('phone')
        email = request.form.get('email')
        workplace_name = request.form.get('workplace_name')
        workplace_type = request.form.get('workplace_type')
        affiliation_eng = request.form.get('affiliation_eng')
        rating_required = request.form.get('rating_required')
        registration_fee = request.form.get('registration_fee')
        remark_user = request.form.get('remark_user')
        agree_terms = request.form.get('agree_terms')
        
        # 필수 필드 검증
        required_fields = {
            'registration_type': registration_type,
            'name_kor': name_kor,
            'first_name': first_name,
            'family_name': family_name,
            'license_number': license_number,
            'specialty': specialty,
            'phone': phone,
            'email': email,
            'workplace_name': workplace_name,
            'workplace_type': workplace_type,
            'affiliation_eng': affiliation_eng,
            'rating_required': rating_required,
            'registration_fee': registration_fee
        }
        
        for field, value in required_fields.items():
            if not value or not value.strip():
                return jsonify({
                    'success': False,
                    'message': f'{field} 필드는 필수입니다.'
                }), 400
        
        # 약관 동의 확인
        if not agree_terms:
            return jsonify({
                'success': False,
                'message': '개인정보 처리방침에 동의해주세요.'
            }), 400
        
        # 이메일 중복 확인
        existing_email = Participant.query.filter_by(email=email, event_id=event_id).first()
        if existing_email:
            return jsonify({
                'success': False,
                'message': '이미 등록된 이메일입니다.'
            }), 400
        
        # 면허번호 중복 확인
        existing_license = Participant.query.filter_by(license_number=license_number, event_id=event_id).first()
        if existing_license:
            return jsonify({
                'success': False,
                'message': '이미 등록된 면허번호입니다.'
            }), 400
        
        # 다음 코드 번호 생성
        last_participant = Participant.query.filter_by(event_id=event_id).order_by(Participant.code.desc()).first()
        next_code = (last_participant.code + 1) if last_participant and last_participant.code else 1
        
        # 진료과목 처리
        final_specialty = specialty_other if specialty == '기타' and specialty_other else specialty
        
        # 새 참가자 생성
        new_participant = Participant(
            event_id=event_id,
            code=next_code,
            registration=registration_type,
            division=final_specialty,
            role=registration_type,
            name_kor=name_kor,
            first_name=first_name,
            family_name=family_name,
            affiliation_eng=affiliation_eng,
            email=email,
            phone=phone,
            license_number=license_number,
            position=workplace_type,
            agree=agree_terms,
            remark_user=remark_user,
            accept_or_decline='Accept'  # 자동으로 수락 상태로 설정
        )
        
        # 데이터베이스에 저장
        db.session.add(new_participant)
        db.session.commit()
        
        print(f"새 참가자 등록 완료: {name_kor} ({email}) - 코드: {next_code}")
        
        return jsonify({
            'success': True,
            'message': '등록이 완료되었습니다.',
            'registration_id': next_code,
            'name_kor': name_kor,
            'email': email
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"참가자 등록 오류: {e}")
        return jsonify({
            'success': False,
            'message': f'등록 중 오류가 발생했습니다: {str(e)}'
        }), 500

# 면허번호 중복 확인 API
@app.route('/check_license_duplicate', methods=['POST'])
def check_license_duplicate():
    """면허번호 중복 확인"""
    try:
        data = request.get_json()
        license_number = data.get('license_number')
        
        if not license_number:
            return jsonify({'exists': False, 'message': '면허번호를 입력해주세요.'}), 400
        
        # 모든 이벤트에서 면허번호 중복 확인
        existing = Participant.query.filter_by(license_number=license_number).first()
        
        return jsonify({
            'exists': existing is not None,
            'message': '이미 등록된 면허번호입니다.' if existing else '사용 가능한 면허번호입니다.'
        })
        
    except Exception as e:
        print(f"면허번호 중복 확인 오류: {e}")
        return jsonify({'exists': False, 'message': '오류가 발생했습니다.'}), 500

# 이메일 중복 확인 API
@app.route('/check_email_duplicate', methods=['POST'])
def check_email_duplicate():
    """이메일 중복 확인"""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'exists': False, 'message': '이메일을 입력해주세요.'}), 400
        
        # 모든 이벤트에서 이메일 중복 확인
        existing = Participant.query.filter_by(email=email).first()
        
        return jsonify({
            'exists': existing is not None,
            'message': '이미 등록된 이메일입니다.' if existing else '사용 가능한 이메일입니다.'
        })
        
    except Exception as e:
        print(f"이메일 중복 확인 오류: {e}")
        return jsonify({'exists': False, 'message': '오류가 발생했습니다.'}), 500

# 행사 장소 업데이트 API
@app.route('/update_event_location/<int:event_id>', methods=['POST'])
def update_event_location(event_id):
    """행사 장소 업데이트"""
    try:
        print(f"=== 장소 업데이트 요청 ===")
        print(f"Event ID: {event_id}")
        print(f"Request data: {request.get_json()}")
        
        data = request.get_json()
        location = data.get('location')
        
        print(f"Location: {location}")
        
        if not location:
            print("장소가 비어있음")
            return jsonify({'success': False, 'message': '장소를 입력해주세요.'}), 400
        
        event = Event.query.get_or_404(event_id)
        print(f"Found event: {event.name}")
        print(f"Current location: {event.location}")
        
        event.location = location
        db.session.commit()
        
        print(f"행사 {event_id}의 장소가 '{location}'로 업데이트되었습니다.")
        
        return jsonify({'success': True, 'message': '장소가 성공적으로 업데이트되었습니다.'})
        
    except Exception as e:
        db.session.rollback()
        print(f"장소 업데이트 오류: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'장소 업데이트 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/download_file_path/<path:filepath>', methods=['GET'])
def download_file_path(filepath):
    """파일 경로로 다운로드"""
    try:
        return send_file(filepath, as_attachment=True)
    except Exception as e:
        logging.error(f"File download failed: {e}")
        return "File not found", 404

@app.route('/get_participants_for_qr/<int:event_id>', methods=['POST'])
def get_participants_for_qr(event_id):
    """QR 코드 생성을 위한 참가자 데이터 JSON 반환"""
    event = Event.query.get_or_404(event_id)
    participant_ids = request.form.get('selected_participants', '')
    if participant_ids:
        ids = [int(pid) for pid in participant_ids.split(',') if pid.strip().isdigit()]
        participants = Participant.query.filter(Participant.id.in_(ids)).all()
    else:
        participants = Participant.query.filter_by(event_id=event_id).all()
    
    # 코드 번호로 오름차순 정렬
    participants = sorted(participants, key=lambda p: int(p.code) if p.code and p.code.isdigit() else 0)
    
    # 행사명에서 안전한 폴더명 생성
    safe_event_name = "".join(c for c in event.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    
    # 데이터 준비
    data = []
    for p in participants:
        data.append({
            'event_id': event.event_id,
            'code': p.code,
            'name_kor': p.name_kor,
            'first_name': p.first_name,
            'family_name': p.family_name,
            'affiliation_kor': p.affiliation_kor,
            'affiliation_eng': p.affiliation_eng,
            'email': p.email,
            'accept_or_decline': p.accept_or_decline
        })
    
    return jsonify({
        'participants': data,
        'event_name': safe_event_name
    })

@app.route('/generate_qr_image/<int:code>/<path:event_name>')
def generate_qr_image(code, event_name):
    """QR 코드 이미지 생성 및 반환"""
    try:
        # QR 코드 생성
        qr = qrcode.make(str(code))
        
        # 이미지를 바이트 스트림으로 변환
        img_io = BytesIO()
        qr.save(img_io, 'PNG')
        img_io.seek(0)
        
        return send_file(
            img_io,
            mimetype='image/png',
            as_attachment=True,
            download_name=f'QR_Codes/{event_name}/{code}.png'
        )
    except Exception as e:
        logging.error(f"QR generation failed for {code}: {str(e)}")
        return "QR generation failed", 500

@app.route('/download_qr_participants_excel/<int:event_id>', methods=['POST'])
def download_qr_participants_excel(event_id):
    event = Event.query.get_or_404(event_id)
    participant_ids = request.form.get('selected_participants', '')
    if participant_ids:
        ids = [int(pid) for pid in participant_ids.split(',') if pid.strip().isdigit()]
        participants = Participant.query.filter(Participant.id.in_(ids)).order_by(Participant.code).all()
    else:
        participants = Participant.query.filter_by(event_id=event_id).order_by(Participant.code).all()

    import tempfile, shutil, zipfile
    temp_dir = tempfile.mkdtemp()
    qr_dir = os.path.join(temp_dir, "QR_Codes")
    os.makedirs(qr_dir, exist_ok=True)

    # QR 코드 생성 및 경로 수집
    data = []
    
    # 코드 번호로 오름차순 정렬 (모든 방식에 공통 적용)
    def safe_code_sort_key(p):
        if not p.code:
            return 999999  # 코드가 없는 경우 맨 뒤로
        try:
            # 문자열을 정수로 변환하여 정렬
            return int(str(p.code))
        except (ValueError, TypeError):
            return 999999  # 변환 실패 시 맨 뒤로
    
    # 정렬 전 상태 로깅
    print(f"DEBUG: Before sorting - participants codes: {[p.code for p in participants]}")
    
    participants = sorted(participants, key=safe_code_sort_key)
    
    # 정렬 후 상태 로깅
    print(f"DEBUG: After sorting - participants codes: {[p.code for p in participants]}")
    
    # 디버깅: 정렬 결과 확인
    print(f"DEBUG: Sorted participants codes: {[p.code for p in participants]}")
    
    export_layout = request.form.get('export_layout', 'normal')
    if export_layout in ('layout_2', 'layout_3', 'layout_4'):
        # 한 행에 N명 배치 (_1, _2, _3, _4)
        import math
        per_row = int(export_layout.split('_')[1])
        max_rows = math.ceil(len(participants) / per_row) if participants else 0
        for i in range(max_rows):
            row_data = {}
            for j in range(per_row):
                idx = i * per_row + j
                suffix = f"_{j + 1}"
                if idx < len(participants):
                    p = participants[idx]
                    if p.code:
                        qr_filename = f"{p.code}.png"
                        qr_path = os.path.join(qr_dir, qr_filename)
                        qr = qrcode.make(str(p.code))
                        qr.save(qr_path)
                        image_path_for_excel = f"QR_Codes/{p.code}.png"
                    else:
                        image_path_for_excel = ""
                    row_data.update({
                        f'Event ID{suffix}': event.event_id,
                        f'Code{suffix}': p.code,
                        f'@Image{suffix}': image_path_for_excel,
                        f'Name (KOR){suffix}': p.name_kor,
                        f'Name (ENG){suffix}': f"{p.first_name} {p.family_name}".strip(),
                        f'Affiliation{suffix}': p.affiliation_kor or p.affiliation_eng,
                        f'Email{suffix}': p.email,
                        f'Accept/Decline{suffix}': p.accept_or_decline
                    })
                else:
                    row_data.update({
                        f'Event ID{suffix}': '',
                        f'Code{suffix}': '',
                        f'@Image{suffix}': '',
                        f'Name (KOR){suffix}': '',
                        f'Name (ENG){suffix}': '',
                        f'Affiliation{suffix}': '',
                        f'Email{suffix}': '',
                        f'Accept/Decline{suffix}': ''
                    })
            data.append(row_data)
    else:
        # 일반 방식 (코드번호 순서대로)
        for p in participants:
            if p.code:
                qr_filename = f"{p.code}.png"
                qr_path = os.path.join(qr_dir, qr_filename)
                qr = qrcode.make(str(p.code))
                qr.save(qr_path)
                image_path_for_excel = f"QR_Codes/{p.code}.png"
            else:
                image_path_for_excel = ""
            data.append({
                'Event ID': event.event_id,
                'Code': p.code,
                '@Image': image_path_for_excel,
                'Name (KOR)': p.name_kor,
                'Name (ENG)': f"{p.first_name} {p.family_name}".strip(),
                'Affiliation': p.affiliation_kor or p.affiliation_eng,
                'Email': p.email,
                'Accept/Decline': p.accept_or_decline
            })

    # Excel, TXT 파일 생성
    import pandas as pd
    df = pd.DataFrame(data)
    excel_path = os.path.join(temp_dir, "participants.xlsx")
    df.to_excel(excel_path, sheet_name='Participants', index=False)
    txt_path = os.path.join(temp_dir, "participants.txt")
    notice = "※ 압축을 푼 폴더에서 QR_Codes 폴더와 participants.xlsx, participants.txt를 사용하세요. @Image 컬럼은 QR_Codes 폴더의 이미지를 가리킵니다.\n"
    with open(txt_path, 'w', encoding='utf-16') as f:
        f.write(notice)
        df.to_csv(f, sep='\t', index=False, encoding='utf-16', lineterminator='\n')

    # ZIP 파일 생성 (Excel_QR_Code/ 루트 폴더로 감싸기)
    zip_path = os.path.join(temp_dir, "Excel_QR_Code.zip")
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        zipf.write(excel_path, "Excel_QR_Code/participants.xlsx")
        zipf.write(txt_path, "Excel_QR_Code/participants.txt")
        for filename in os.listdir(qr_dir):
            zipf.write(os.path.join(qr_dir, filename), f"Excel_QR_Code/QR_Codes/{filename}")

    from flask import after_this_request
    @after_this_request
    def cleanup(response):
        shutil.rmtree(temp_dir)
        return response

    return send_file(
        zip_path,
        mimetype='application/zip',
        as_attachment=True,
        download_name='Excel_QR_Code.zip'
    )

@app.route('/download_qr_participants_excel_with_path/<int:event_id>', methods=['POST'])
def download_qr_participants_excel_with_path(event_id):
    event = Event.query.get_or_404(event_id)
    participant_ids = request.form.get('selected_participants', '')
    custom_path = request.form.get('custom_path', '/Users/jhc/Downloads/Excel_QR_Code')
    if participant_ids:
        ids = [int(pid) for pid in participant_ids.split(',') if pid.strip().isdigit()]
        participants = Participant.query.filter(Participant.id.in_(ids)).order_by(Participant.code).all()
    else:
        participants = Participant.query.filter_by(event_id=event_id).order_by(Participant.code).all()
    
    # 코드 번호로 오름차순 정렬 (한 번만)
    def safe_code_sort_key(p):
        if not p.code:
            return 999999  # 코드가 없는 경우 맨 뒤로
        try:
            # 문자열을 정수로 변환하여 정렬
            return int(str(p.code))
        except (ValueError, TypeError):
            return 999999  # 변환 실패 시 맨 뒤로
    
    # 정렬 전 상태 로깅
    print(f"DEBUG: Before sorting - participants codes: {[p.code for p in participants]}")
    
    participants = sorted(participants, key=safe_code_sort_key)
    
    # 정렬 후 상태 로깅
    print(f"DEBUG: After sorting - participants codes: {[p.code for p in participants]}")
    
    # 경로 정규화 (백슬래시 이스케이프 처리)
    custom_path = custom_path.replace('\\', '/')
    
    # QR 코드 생성 및 실제 경로 수집
    data = []
    
    export_layout = request.form.get('export_layout', 'normal')
    if export_layout in ('layout_2', 'layout_3', 'layout_4'):
        # 한 행에 N명 배치 (_1, _2, _3, _4)
        import math
        per_row = int(export_layout.split('_')[1])
        max_rows = math.ceil(len(participants) / per_row) if participants else 0
        for i in range(max_rows):
            row_data = {}
            for j in range(per_row):
                idx = i * per_row + j
                suffix = f"_{j + 1}"
                if idx < len(participants):
                    p = participants[idx]
                    if p.code:
                        actual_image_path = f"{custom_path}/Excel_QR_Code/QR_Codes/{p.code}.png"
                    else:
                        actual_image_path = ""
                    row_data.update({
                        f'Event ID{suffix}': event.event_id,
                        f'Code{suffix}': p.code,
                        f'@Image{suffix}': actual_image_path,
                        f'Name (KOR){suffix}': p.name_kor,
                        f'Name (ENG){suffix}': f"{p.first_name} {p.family_name}".strip(),
                        f'Affiliation{suffix}': p.affiliation_kor or p.affiliation_eng,
                        f'Email{suffix}': p.email,
                        f'Accept/Decline{suffix}': p.accept_or_decline
                    })
                else:
                    row_data.update({
                        f'Event ID{suffix}': '',
                        f'Code{suffix}': '',
                        f'@Image{suffix}': '',
                        f'Name (KOR){suffix}': '',
                        f'Name (ENG){suffix}': '',
                        f'Affiliation{suffix}': '',
                        f'Email{suffix}': '',
                        f'Accept/Decline{suffix}': ''
                    })
            data.append(row_data)
    else:
        # 일반 방식 (순서대로, 이미 정렬되어 있음)
        for p in participants:
            if p.code:
                actual_image_path = f"{custom_path}/Excel_QR_Code/QR_Codes/{p.code}.png"
            else:
                actual_image_path = ""
            data.append({
                'Event ID': event.event_id,
                'Code': p.code,
                '@Image': actual_image_path,
                'Name (KOR)': p.name_kor,
                'Name (ENG)': f"{p.first_name} {p.family_name}".strip(),
                'Affiliation': p.affiliation_kor or p.affiliation_eng,
                'Email': p.email,
                'Accept/Decline': p.accept_or_decline
            })

    # Excel 파일 생성
    import pandas as pd
    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Participants')
        worksheet = writer.sheets['Participants']
        
        # 컬럼 너비 자동 조정
        for col_num, value in enumerate(df.columns.values):
            max_length = max(
                df[value].astype(str).apply(len).max(),
                len(str(value))
            )
            worksheet.set_column(col_num, col_num, min(max_length + 2, 50))
    
    output.seek(0)
    current_date = datetime.now().strftime("%Y%m%d")
    filename = f"{event.event_id}_participants_with_paths_{current_date}.xlsx"
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )

@app.route('/download_qr_participants_excel_with_path_zip/<int:event_id>', methods=['POST'])
def download_qr_participants_excel_with_path_zip(event_id):
    event = Event.query.get_or_404(event_id)
    participant_ids = request.form.get('selected_participants', '')
    custom_path = request.form.get('custom_path', '/Users/jhc/Downloads/Excel_QR_Code')
    if participant_ids:
        ids = [int(pid) for pid in participant_ids.split(',') if pid.strip().isdigit()]
        participants = Participant.query.filter(Participant.id.in_(ids)).order_by(Participant.code).all()
    else:
        participants = Participant.query.filter_by(event_id=event_id).order_by(Participant.code).all()

    import tempfile, shutil, zipfile
    temp_dir = tempfile.mkdtemp()
    qr_dir = os.path.join(temp_dir, "QR_Codes")
    os.makedirs(qr_dir, exist_ok=True)

    # 경로 정규화 (백슬래시 이스케이프 처리)
    custom_path = custom_path.replace('\\', '/')
    
    # QR 코드 생성 및 실제 경로 수집
    data = []
    
    # 코드 번호로 오름차순 정렬 (모든 방식에 공통 적용)
    def safe_code_sort_key(p):
        if not p.code:
            return 999999  # 코드가 없는 경우 맨 뒤로
        try:
            # 문자열을 정수로 변환하여 정렬
            return int(str(p.code))
        except (ValueError, TypeError):
            return 999999  # 변환 실패 시 맨 뒤로
    
    # 정렬 전 상태 로깅
    print(f"DEBUG: Before sorting - participants codes: {[p.code for p in participants]}")
    
    participants = sorted(participants, key=safe_code_sort_key)
    
    # 정렬 후 상태 로깅
    print(f"DEBUG: After sorting - participants codes: {[p.code for p in participants]}")
    
    export_layout = request.form.get('export_layout', 'normal')
    if export_layout in ('layout_2', 'layout_3', 'layout_4'):
        # 한 행에 N명 배치 (_1, _2, _3, _4)
        import math
        per_row = int(export_layout.split('_')[1])
        max_rows = math.ceil(len(participants) / per_row) if participants else 0
        for i in range(max_rows):
            row_data = {}
            for j in range(per_row):
                idx = i * per_row + j
                suffix = f"_{j + 1}"
                if idx < len(participants):
                    p = participants[idx]
                    if p.code:
                        qr_filename = f"{p.code}.png"
                        qr_path = os.path.join(qr_dir, qr_filename)
                        qr = qrcode.make(str(p.code))
                        qr.save(qr_path)
                        actual_image_path = f"{custom_path}/Excel_QR_Code/QR_Codes/{p.code}.png"
                    else:
                        actual_image_path = ""
                    row_data.update({
                        f'Event ID{suffix}': event.event_id,
                        f'Code{suffix}': p.code,
                        f'@Image{suffix}': actual_image_path,
                        f'Name (KOR){suffix}': p.name_kor,
                        f'Name (ENG){suffix}': f"{p.first_name} {p.family_name}".strip(),
                        f'Affiliation{suffix}': p.affiliation_kor or p.affiliation_eng,
                        f'Email{suffix}': p.email,
                        f'Accept/Decline{suffix}': p.accept_or_decline
                    })
                else:
                    row_data.update({
                        f'Event ID{suffix}': '',
                        f'Code{suffix}': '',
                        f'@Image{suffix}': '',
                        f'Name (KOR){suffix}': '',
                        f'Name (ENG){suffix}': '',
                        f'Affiliation{suffix}': '',
                        f'Email{suffix}': '',
                        f'Accept/Decline{suffix}': ''
                    })
            data.append(row_data)
    else:
        # 일반 방식 (코드번호 순서대로)
        for p in participants:
            if p.code:
                qr_filename = f"{p.code}.png"
                qr_path = os.path.join(qr_dir, qr_filename)
                qr = qrcode.make(str(p.code))
                qr.save(qr_path)
                actual_image_path = f"{custom_path}/Excel_QR_Code/QR_Codes/{p.code}.png"
            else:
                actual_image_path = ""
            data.append({
                'Event ID': event.event_id,
                'Code': p.code,
                '@Image': actual_image_path,
                'Name (KOR)': p.name_kor,
                'Name (ENG)': f"{p.first_name} {p.family_name}".strip(),
                'Affiliation': p.affiliation_kor or p.affiliation_eng,
                'Email': p.email,
                'Accept/Decline': p.accept_or_decline
            })

    # Excel, TXT 파일 생성
    import pandas as pd
    df = pd.DataFrame(data)
    excel_path = os.path.join(temp_dir, "participants.xlsx")
    df.to_excel(excel_path, sheet_name='Participants', index=False)
    
    txt_path = os.path.join(temp_dir, "participants.txt")
    with open(txt_path, 'w', encoding='utf-16') as f:
        df.to_csv(f, sep='\t', index=False, encoding='utf-16', lineterminator='\n')

    # ZIP 파일 생성 (Excel_QR_Code/ 루트 폴더로 감싸기)
    zip_path = os.path.join(temp_dir, "Excel_QR_Code.zip")
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        zipf.write(excel_path, "Excel_QR_Code/participants.xlsx")
        zipf.write(txt_path, "Excel_QR_Code/participants.txt")
        for filename in os.listdir(qr_dir):
            zipf.write(os.path.join(qr_dir, filename), f"Excel_QR_Code/QR_Codes/{filename}")

    # 사용자 지정 경로에 직접 저장 (localhost이므로 서버가 파일시스템 접근 가능)
    # Excel @Image: {custom_path}/Excel_QR_Code/QR_Codes/{code}.png
    save_path = None
    is_absolute = custom_path and (custom_path.startswith('/') or (len(custom_path) > 1 and custom_path[1] == ':'))
    if is_absolute:
        try:
            save_dir = os.path.normpath(custom_path.rstrip('/\\'))
            os.makedirs(save_dir, exist_ok=True)
            dest_zip = os.path.join(save_dir, 'Excel_QR_Code.zip')
            shutil.copy2(zip_path, dest_zip)
            save_path = dest_zip
        except Exception as e:
            logging.warning(f"Could not save ZIP to {custom_path}: {e}")

    from flask import after_this_request
    @after_this_request
    def cleanup(response):
        shutil.rmtree(temp_dir)
        return response

    if save_path:
        return jsonify({'success': True, 'saved_to': save_path})
    return send_file(
        zip_path,
        mimetype='application/zip',
        as_attachment=True,
        download_name='Excel_QR_Code.zip'
    )

@app.route('/download_custom_excel/<int:event_id>', methods=['POST'])
def download_custom_excel(event_id):
    """평점용 서식으로 엑셀 다운로드"""
    event = Event.query.get_or_404(event_id)
    participant_ids = request.form.get('selected_participants', '')
    rating_criteria = request.form.get('rating_criteria', '6')
    
    if participant_ids:
        ids = [int(pid) for pid in participant_ids.split(',') if pid.strip().isdigit()]
        participants = Participant.query.filter(Participant.id.in_(ids)).all()
    else:
        participants = Participant.query.filter_by(event_id=event_id).all()

    # Code 번호로 오름차순 정렬 (안전한 처리)
    def safe_code_sort_key(p):
        try:
            if p.code and str(p.code).isdigit():
                return int(p.code)
            return 0
        except (ValueError, TypeError):
            return 0
    
    participants = sorted(participants, key=safe_code_sort_key)

    data = []
    for p in participants:
        # 체크인/아웃 시간 처리
        check_in_dt = pd.to_datetime(p.check_in_time, errors='coerce') if p.check_in_time else None
        check_out_dt = pd.to_datetime(p.check_out_time, errors='coerce') if p.check_out_time else None
        
        check_in_time_str = check_in_dt.strftime('%H:%M') if pd.notna(check_in_dt) else ''
        check_out_time_str = check_out_dt.strftime('%H:%M') if pd.notna(check_out_dt) else ''
        
        check_in_status = 1 if check_in_time_str else 0
        check_out_status = 1 if check_out_time_str else 0
        
        # 체류시간 계산
        duration = (check_out_dt - check_in_dt) if (pd.notna(check_in_dt) and pd.notna(check_out_dt)) else None
        duration_seconds = duration.total_seconds() if duration and duration.total_seconds() >= 0 else 0
        duration_str = f"{int(duration_seconds // 3600)}시간 {int((duration_seconds % 3600) // 60)}분" if duration and duration.total_seconds() >= 0 else ''
        
        # 평점 계산 (올바른 로직)
        if check_in_status == 0 or check_out_status == 0:
            rating = 0
        else:
            hours = duration_seconds / 3600  # 시간 단위로 변환
            
            if rating_criteria == '1':
                # 1시간 미만=0점, 1시간 이상=1점
                rating = 1 if hours >= 1 else 0
            elif rating_criteria == '2':
                # 1시간 미만=0점, 1-2시간=1점, 2시간 이상=2점
                if hours < 1:
                    rating = 0
                elif hours < 2:
                    rating = 1
                else:
                    rating = 2
            elif rating_criteria == '3':
                # 1시간 미만=0점, 1-2시간=1점, 2-3시간=2점, 3시간 이상=3점
                if hours < 1:
                    rating = 0
                elif hours < 2:
                    rating = 1
                elif hours < 3:
                    rating = 2
                else:
                    rating = 3
            elif rating_criteria == '4':
                # 1시간 미만=0점, 1-2시간=1점, 2-3시간=2점, 3-4시간=3점, 4시간 이상=4점
                if hours < 1:
                    rating = 0
                elif hours < 2:
                    rating = 1
                elif hours < 3:
                    rating = 2
                elif hours < 4:
                    rating = 3
                else:
                    rating = 4
            elif rating_criteria == '5':
                # 1시간 미만=0점, 1-2시간=1점, 2-3시간=2점, 3-4시간=3점, 4-5시간=4점, 5시간 이상=5점
                if hours < 1:
                    rating = 0
                elif hours < 2:
                    rating = 1
                elif hours < 3:
                    rating = 2
                elif hours < 4:
                    rating = 3
                elif hours < 5:
                    rating = 4
                else:
                    rating = 5
            else:  # '6'
                # 1시간 미만=0점, 1-2시간=1점, 2-3시간=2점, 3-4시간=3점, 4-5시간=4점, 5-6시간=5점, 6시간 이상=6점
                if hours < 1:
                    rating = 0
                elif hours < 2:
                    rating = 1
                elif hours < 3:
                    rating = 2
                elif hours < 4:
                    rating = 3
                elif hours < 5:
                    rating = 4
                elif hours < 6:
                    rating = 5
                else:
                    rating = 6
        
        data.append({
            '연번': p.code if p.code else '',  # 참가자의 Code 번호 사용 (None 처리)
            '성명': p.name_kor or f'{p.first_name} {p.family_name}',
            '의사 면허번호': p.license_number,
            '교육전 서명 시간': check_in_time_str,
            '서명여부 (1=yes, 0=no) (Check-in)': check_in_status,
            '교육후 서명 시간': check_out_time_str,
            '서명여부 (1=yes, 0=no) (Check-out)': check_out_status,
            '체류시간 (자동계산)': duration_str,
            '발급평점': rating,
            '소속': p.affiliation_kor or p.affiliation_eng,
            '직업구분': p.division,
            '비고1': p.department_kor or p.department_eng
        })

    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name=f'{rating_criteria}평점')
        worksheet = writer.sheets[f'{rating_criteria}평점']
        for col_num, _ in enumerate(df.columns):
            worksheet.set_column(col_num, col_num, None, writer.book.add_format({'num_format': '@'}))
    
    output.seek(0)
    current_date = datetime.now().strftime("%Y%m%d")
    filename = f"{event.event_id}_평점신고_{rating_criteria}평점_{current_date}.xlsx"
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )

# 행사 프로그램 관리 라우트들
@app.route('/event_program/<int:event_id>')
def event_program(event_id):
    """행사 프로그램 관리 페이지"""
    try:
        event = Event.query.get(event_id)
        if not event:
            return "행사를 찾을 수 없습니다.", 404
        
        return render_template('event_program.html', event=event)
    except Exception as e:
        logging.error(f"Error loading event program page: {str(e)}")
        return "페이지 로드 중 오류가 발생했습니다.", 500

@app.route('/api/event_program/<int:event_id>', methods=['GET'])
def get_event_program(event_id):
    """행사 프로그램 데이터 조회"""
    try:
        # 임시로 JSON 파일에서 데이터를 읽어오는 방식
        # 실제로는 데이터베이스 테이블을 만들어야 함
        program_file = f'uploads/event_program_{event_id}.json'
        
        if os.path.exists(program_file):
            with open(program_file, 'r', encoding='utf-8') as f:
                import json
                data = json.load(f)
                return jsonify(data)
        else:
            return jsonify({'sessions': []})
    except Exception as e:
        logging.error(f"Error getting event program: {str(e)}")
        return jsonify({'error': '프로그램 데이터 조회 중 오류가 발생했습니다.'}), 500

@app.route('/api/event_program/<int:event_id>', methods=['POST'])
def save_event_program(event_id):
    """행사 프로그램 데이터 저장"""
    try:
        data = request.get_json()
        sessions = data.get('sessions', [])
        
        # JSON 파일로 저장 (임시 방식)
        # 실제로는 데이터베이스 테이블을 만들어야 함
        program_file = f'uploads/event_program_{event_id}.json'
        os.makedirs('uploads', exist_ok=True)
        
        with open(program_file, 'w', encoding='utf-8') as f:
            import json
            json.dump({'sessions': sessions}, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True})
    except Exception as e:
        logging.error(f"Error saving event program: {str(e)}")
        return jsonify({'error': '프로그램 저장 중 오류가 발생했습니다.'}), 500

@app.route('/api/event_program/<int:event_id>/export')
def export_event_program(event_id):
    """행사 프로그램 Excel 내보내기"""
    try:
        event = Event.query.get(event_id)
        if not event:
            return "행사를 찾을 수 없습니다.", 404
        
        # 프로그램 데이터 로드
        program_file = f'uploads/event_program_{event_id}.json'
        if not os.path.exists(program_file):
            return jsonify({'error': '프로그램 데이터가 없습니다.'}), 404
        
        with open(program_file, 'r', encoding='utf-8') as f:
            import json
            data = json.load(f)
            sessions = data.get('sessions', [])
        
        # Excel 데이터 준비
        excel_data = []
        for session in sessions:
            # 세션 정보
            session_row = {
                '시간': f"{session['startTime']} - {session['endTime']}",
                '세션': session['title'],
                '좌장': session['chair'],
                '발표자': '',
                '주제': '',
                '발표시간': ''
            }
            excel_data.append(session_row)
            
            # 발표자 정보
            for speaker in session.get('speakers', []):
                speaker_row = {
                    '시간': '',
                    '세션': '',
                    '좌장': '',
                    '발표자': speaker['name'],
                    '주제': speaker['topic'],
                    '발표시간': f"{speaker['startTime']} - {speaker['endTime']}"
                }
                excel_data.append(speaker_row)
            
            # 빈 줄 추가
            excel_data.append({
                '시간': '',
                '세션': '',
                '좌장': '',
                '발표자': '',
                '주제': '',
                '발표시간': ''
            })
        
        df = pd.DataFrame(excel_data)
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='행사프로그램')
            worksheet = writer.sheets['행사프로그램']
            
            # 열 너비 조정
            worksheet.set_column('A:A', 15)  # 시간
            worksheet.set_column('B:B', 25)  # 세션
            worksheet.set_column('C:C', 20)  # 좌장
            worksheet.set_column('D:D', 20)  # 발표자
            worksheet.set_column('E:E', 30)  # 주제
            worksheet.set_column('F:F', 15)  # 발표시간
        
        output.seek(0)
        filename = f"event_program_{event_id}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logging.error(f"Error exporting event program: {str(e)}")
        return jsonify({'error': '프로그램 내보내기 중 오류가 발생했습니다.'}), 500

# ===== 강의 장소 관리 API =====

@app.route('/api/event_program/<int:event_id>/venues', methods=['GET'])
def get_event_venues(event_id):
    """행사 강의 장소 목록 조회"""
    try:
        event = Event.query.get(event_id)
        if not event:
            return jsonify({'error': '행사를 찾을 수 없습니다.'}), 404
        
        # JSON 파일에서 강의 장소 데이터 로드
        venues_file = f'uploads/event_venues_{event_id}.json'
        if os.path.exists(venues_file):
            with open(venues_file, 'r', encoding='utf-8') as f:
                import json
                data = json.load(f)
                venues = data.get('venues', [])
        else:
            venues = []
        
        return jsonify({'success': True, 'venues': venues})
    except Exception as e:
        logging.error(f"Error loading event venues: {str(e)}")
        return jsonify({'error': '강의 장소 로드 중 오류가 발생했습니다.'}), 500

@app.route('/api/event_program/<int:event_id>/venues', methods=['POST'])
def save_event_venues(event_id):
    """행사 강의 장소 저장"""
    try:
        event = Event.query.get(event_id)
        if not event:
            return jsonify({'error': '행사를 찾을 수 없습니다.'}), 404
        
        data = request.get_json()
        venues = data.get('venues', [])
        
        # JSON 파일로 저장
        venues_file = f'uploads/event_venues_{event_id}.json'
        os.makedirs('uploads', exist_ok=True)
        
        with open(venues_file, 'w', encoding='utf-8') as f:
            import json
            json.dump({'venues': venues}, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True})
    except Exception as e:
        logging.error(f"Error saving event venues: {str(e)}")
        return jsonify({'error': '강의 장소 저장 중 오류가 발생했습니다.'}), 500

@app.route('/api/event_program/<int:event_id>/participants', methods=['GET'])
def get_event_participants(event_id):
    """행사 참가자 목록 조회 (좌장/발표자 선택용)"""
    try:
        event = Event.query.get(event_id)
        if not event:
            return jsonify({'error': '행사를 찾을 수 없습니다.'}), 404
        
        participants = Participant.query.filter_by(event_id=event_id).all()
        
        participants_list = []
        for p in participants:
            # 참가자 이름 생성 (한글명 우선, 없으면 영문명)
            display_name = p.name_kor if p.name_kor else (p.first_name + ' ' + p.family_name if p.first_name and p.family_name else p.name_eng or '참가자')
            
            participants_list.append({
                'id': p.id,
                'name': display_name,
                'name_kor': p.name_kor,
                'name_eng': p.name_eng,
                'first_name': p.first_name,
                'family_name': p.family_name,
                'email': p.email,
                'phone': p.phone,
                'affiliation_kor': p.affiliation_kor,
                'affiliation_eng': p.affiliation_eng,
                'department_kor': p.department_kor,
                'department_eng': p.department_eng,
                'country': p.country,
                'country_code': p.country_code,
                'role': p.role,
                'position': p.position
            })
        
        return jsonify({'success': True, 'participants': participants_list})
    except Exception as e:
        logging.error(f"Error loading event participants: {str(e)}")
        return jsonify({'error': '참가자 목록 로드 중 오류가 발생했습니다.'}), 500

@app.route('/api/event/<int:event_id>/participants', methods=['POST'])
def add_event_participant(event_id):
    """행사 참가자 빠른 추가 (엑셀 업로드 중)"""
    try:
        event = Event.query.get(event_id)
        if not event:
            return jsonify({'error': '행사를 찾을 수 없습니다.'}), 404
        
        data = request.get_json()
        
        # 필수 필드 확인
        name = data.get('name')
        email = data.get('email')
        
        if not name or not email:
            return jsonify({'error': '이름과 이메일은 필수입니다.'}), 400
        
        # 이메일 중복 확인
        existing = Participant.query.filter_by(event_id=event_id, email=email).first()
        if existing:
            return jsonify({'error': '이미 등록된 이메일입니다.'}), 400
        
        # 새 참가자 코드 생성
        max_code = db.session.query(db.func.max(Participant.code)).filter_by(event_id=event_id).scalar()
        new_code = str(int(max_code) + 1 if max_code else 1)
        
        # 참가자 생성
        participant = Participant(
            event_id=event_id,
            code=new_code,
            name_kor=name,
            email=email,
            affiliation_kor=data.get('affiliation', ''),
            name_eng=data.get('name_eng', ''),
            phone=data.get('phone', ''),
            registration=data.get('registration', 'Online'),  # 실제 컬럼
            accept_or_decline=data.get('accept_or_decline', 'Accept')  # 실제 컬럼
        )
        
        db.session.add(participant)
        db.session.commit()
        
        # 응답 데이터 준비
        participant_data = {
            'id': participant.id,
            'name': participant.name_kor,
            'email': participant.email,
            'affiliation': participant.affiliation_kor,
            'name_eng': participant.name_eng
        }
        
        return jsonify({'success': True, 'participant': participant_data})
    except Exception as e:
        db.session.rollback()
        import traceback
        error_details = traceback.format_exc()
        logging.error(f"Error adding participant: {str(e)}")
        logging.error(f"Full traceback: {error_details}")
        return jsonify({'error': f'{str(e)}'}), 500

@app.route('/api/event/<int:event_id>/date', methods=['PUT'])
def update_event_date(event_id):
    """행사 날짜 업데이트"""
    try:
        event = Event.query.get(event_id)
        if not event:
            return jsonify({'error': '행사를 찾을 수 없습니다.'}), 404
        
        data = request.get_json()
        field = data.get('field')
        value = data.get('value')
        
        if field not in ['start_date', 'end_date']:
            return jsonify({'error': '잘못된 필드입니다.'}), 400
        
        if not value:
            return jsonify({'error': '날짜 값을 입력해주세요.'}), 400
        
        # 날짜 형식 검증 (YYYY-MM-DD)
        import re
        if not re.match(r'^\d{4}-\d{2}-\d{2}$', value):
            return jsonify({'error': '날짜 형식이 올바르지 않습니다. (YYYY-MM-DD)'}), 400
        
        # 날짜 업데이트
        if field == 'start_date':
            event.start_date = value
        else:
            event.end_date = value
        
        db.session.commit()
        
        return jsonify({'success': True, 'message': '날짜가 업데이트되었습니다.'})
    except Exception as e:
        db.session.rollback()
        import traceback
        error_details = traceback.format_exc()
        logging.error(f"Error updating event date: {str(e)}")
        logging.error(f"Full traceback: {error_details}")
        return jsonify({'error': f'{str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5002) 